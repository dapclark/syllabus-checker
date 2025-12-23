#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Fix Layout Tables Tool
Converts tables used for layout into properly structured headings and paragraphs.
Also marks accessibility issues in remaining data tables with bold red feedback.

Usage:
    python3 fix_layout_tables.py YOUR_SYLLABUS.docx

Output:
    - YOUR_SYLLABUS_backup.docx (original file)
    - YOUR_SYLLABUS_fixed.docx (converted file with feedback)
"""

import argparse
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from typing import List, Tuple, Optional
import shutil


class LayoutTableFixer:
    """Detects and fixes tables used for layout instead of data"""

    def __init__(self, doc_path: str):
        self.doc_path = doc_path
        self.doc = Document(doc_path)
        self.layout_tables_fixed = 0
        self.data_tables_kept = 0

    def is_layout_table(self, table) -> Tuple[bool, str]:
        """
        Determine if a table is used for layout vs. data.
        Returns (is_layout, reason)
        """
        total_cells = len(table.rows) * len(table.columns)
        if total_cells == 0:
            return False, ""

        # Count empty cells
        empty_cells = 0
        cells_with_long_text = 0
        cell_lengths = []

        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if not text:
                    empty_cells += 1
                else:
                    cell_lengths.append(len(text))
                    if len(text) > 500:
                        cells_with_long_text += 1

        empty_ratio = empty_cells / total_cells if total_cells > 0 else 0

        # Heuristics for layout tables
        reasons = []

        # High ratio of empty cells
        if empty_ratio > 0.4:
            reasons.append(f"{empty_cells}/{total_cells} empty cells")

        # Contains cells with very long text (paragraphs, not data)
        if cells_with_long_text > 0:
            reasons.append(f"{cells_with_long_text} cells with >500 chars")

        # High variance in cell content length suggests layout, not data
        if cell_lengths and len(cell_lengths) > 1:
            avg_length = sum(cell_lengths) / len(cell_lengths)
            variance = sum((x - avg_length) ** 2 for x in cell_lengths) / len(cell_lengths)
            if variance > 100000:  # High variance
                reasons.append("inconsistent cell content lengths")

        # Large table size might indicate layout
        if len(table.rows) > 10 or len(table.columns) > 5:
            reasons.append(f"large size: {len(table.rows)} rows × {len(table.columns)} cols")

        # If 2+ reasons, it's likely a layout table
        if len(reasons) >= 2:
            return True, "; ".join(reasons)

        return False, ""

    def is_heading_text(self, text: str, has_bold: bool) -> bool:
        """Determine if text should be a heading"""
        text = text.strip()

        if not text:
            return False

        # Skip if it's very long (likely paragraph text)
        if len(text) > 200:
            return False

        # Skip if ends with sentence punctuation (likely paragraph text)
        if text.endswith(('.', '!', '?', ';')):
            return False

        # Check for heading indicators
        word_count = len(text.split())

        # Bold and short = likely heading
        if has_bold and word_count <= 10:
            return True

        # All caps and short = likely heading
        if text.isupper() and word_count <= 10:
            return True

        # Single line and bold = likely heading
        if has_bold and '\n' not in text and word_count <= 15:
            return True

        return False

    def check_paragraph_for_bold(self, paragraph) -> bool:
        """Check if paragraph has significant bold formatting"""
        if not paragraph.runs:
            return False

        bold_chars = sum(len(run.text) for run in paragraph.runs if run.bold)
        total_chars = len(paragraph.text)

        if total_chars == 0:
            return False

        # If >60% is bold, consider it bold text
        return (bold_chars / total_chars) > 0.6

    def is_manual_list_item(self, text: str, para) -> bool:
        """Check if text is a manual list item (uses bullet characters instead of list style)"""
        text = text.strip()
        if not text:
            return False

        # Check if it starts with common manual bullet characters
        manual_bullets = ['•', '·', '○', '■', '□', '►', '▪', '-', '*', '–', '—']
        starts_with_bullet = any(text.startswith(b) for b in manual_bullets)

        # Check if it's NOT using a proper list style
        has_list_style = para.style.name and 'List' in para.style.name
        has_numbering = self.has_numbering(para) if hasattr(self, 'has_numbering') else False

        return starts_with_bullet and not has_list_style and not has_numbering

    def contains_raw_url(self, text: str) -> bool:
        """Check if text contains a raw URL that should have descriptive link text"""
        import re
        # Look for URLs
        url_pattern = r'https?://[^\s]+|www\.[^\s]+'
        return bool(re.search(url_pattern, text))

    def is_all_caps(self, text: str) -> bool:
        """Check if text is ALL CAPS (accessibility issue)"""
        text = text.strip()
        if not text or len(text) < 4:  # Skip very short text
            return False

        # Check if mostly uppercase letters (ignore numbers and punctuation)
        letters = [c for c in text if c.isalpha()]
        if len(letters) < 3:
            return False

        uppercase_ratio = sum(1 for c in letters if c.isupper()) / len(letters)
        return uppercase_ratio > 0.8  # More than 80% uppercase

    def is_title_text(self, text: str, para, is_first_content: bool) -> bool:
        """Check if text appears to be a document title that needs Title style"""
        text = text.strip()
        if not text:
            return False

        # Only consider first content paragraph as potential title
        if not is_first_content:
            return False

        # Must be bold
        if not self.check_paragraph_for_bold(para):
            return False

        # Should be relatively short (like a title)
        if len(text) > 150:
            return False

        # Should not have sentence-ending punctuation
        if text.endswith(('.', '!', '?', ';')):
            return False

        # Check it's not already styled as Title
        if para.style.name and 'Title' in para.style.name:
            return False

        return True

    def has_non_descriptive_link_text(self, para) -> bool:
        """Check if paragraph contains hyperlinks with non-descriptive text"""
        # Look for hyperlinks in the paragraph's XML
        non_descriptive_phrases = [
            'click here', 'here', 'link', 'read more', 'more', 'this link',
            'click', 'download', 'view', 'see here', 'see more'
        ]

        for run in para.runs:
            # Check if this run is part of a hyperlink
            if run._element.getparent().tag.endswith('}hyperlink'):
                link_text = run.text.lower().strip()
                if link_text in non_descriptive_phrases:
                    return True

        return False

    def is_very_long_paragraph(self, text: str) -> bool:
        """Check if paragraph is too long (>300 words)"""
        text = text.strip()
        if not text:
            return False

        word_count = len(text.split())
        return word_count > 300

    def has_manual_alignment(self, text: str) -> bool:
        """Check if text uses multiple tabs/spaces for alignment (accessibility issue)"""
        import re

        # Check for multiple consecutive spaces (2+ spaces, common in alignment)
        if re.search(r'  +', text):  # 2 or more consecutive spaces
            return True

        # Check for tabs (often used for alignment)
        if '\t' in text:
            return True

        return False

    def convert_table_to_structure(self, table, table_index: int):
        """Convert a layout table to proper headings and paragraphs"""
        new_elements = []
        is_first_content = True  # Track if we've seen content yet (for title detection)
        last_heading_level = 0  # Track last heading level to detect skips

        # Process each cell in reading order
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                # Process cell contents in document order (paragraphs and tables interleaved)
                # This ensures nested tables stay with their associated text
                from docx.table import Table as DocxTable
                from docx.text.paragraph import Paragraph

                for child in cell._element:
                    # Check if this is a paragraph
                    if child.tag.endswith('}p'):
                        para = Paragraph(child, cell)
                        last_heading_level = self._process_paragraph(
                            para, new_elements, is_first_content, last_heading_level
                        )
                        if para.text.strip():
                            is_first_content = False
                    # Check if this is a table
                    elif child.tag.endswith('}tbl'):
                        nested_table = DocxTable(child, cell)
                        new_elements.append(('table', None, None, nested_table))

        return new_elements

    def _process_paragraph(self, para, new_elements, is_first_content, last_heading_level):
        """Process a single paragraph and add to elements list. Returns new heading level."""
        text = para.text.strip()
        if not text:
            return last_heading_level

        # Check if paragraph already has a Heading style - preserve it
        if para.style.name.startswith('Heading'):
            level = int(para.style.name.split()[-1])

            # Check for skipped heading level
            if last_heading_level > 0 and level > last_heading_level + 1:
                new_elements.append(('skipped_heading', level, text, para))
                return level

            new_elements.append(('heading', level, text, para))
            return level

        # Check if already styled as Title - preserve it
        if para.style.name and 'Title' in para.style.name:
            new_elements.append(('title', None, text, para))
            return last_heading_level

        has_bold = self.check_paragraph_for_bold(para)

        # Check for title (first bold content that looks like a title)
        if self.is_title_text(text, para, is_first_content):
            new_elements.append(('flagged_title', None, text, para))
            return last_heading_level

        # Check if this looks like it should be a heading
        if self.is_heading_text(text, has_bold):
            # Flag it for user to style - don't guess the level
            new_elements.append(('flagged_heading', None, text, para))
            return last_heading_level

        # Check for manual list items
        if self.is_manual_list_item(text, para):
            new_elements.append(('flagged_list', None, text, para))
            return last_heading_level

        # Check for raw URLs
        if self.contains_raw_url(text):
            new_elements.append(('flagged_url', None, text, para))
            return last_heading_level

        # Check for ALL CAPS text (accessibility issue)
        if self.is_all_caps(text):
            new_elements.append(('flagged_allcaps', None, text, para))
            return last_heading_level

        # Check for non-descriptive link text
        if self.has_non_descriptive_link_text(para):
            new_elements.append(('flagged_link', None, text, para))
            return last_heading_level

        # Check for manual alignment with tabs/spaces
        if self.has_manual_alignment(text):
            new_elements.append(('flagged_alignment', None, text, para))
            return last_heading_level

        # Check for very long paragraphs
        if self.is_very_long_paragraph(text):
            new_elements.append(('flagged_long', None, text, para))
            return last_heading_level

        # Regular paragraph
        new_elements.append(('paragraph', None, text, para))
        return last_heading_level

    def has_numbering(self, para):
        """Check if paragraph has Word numbering (bullets/numbers)"""
        if para._element.pPr is not None:
            numPr = para._element.pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
            return numPr is not None
        return False

    def copy_numbering(self, source_para, target_para):
        """Copy numbering properties from source to target paragraph"""
        if source_para._element.pPr is not None:
            numPr = source_para._element.pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
            if numPr is not None:
                # Ensure target has pPr element
                if target_para._element.pPr is None:
                    from docx.oxml import parse_xml
                    target_para._element.insert(0, parse_xml('<w:pPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))

                # Copy the numbering properties
                import copy
                target_para._element.pPr.append(copy.deepcopy(numPr))

    def insert_elements_before_table(self, table, elements: List[Tuple]):
        """Insert new structured content before the table"""
        # Find the table's position in the document
        table_element = table._element
        parent = table_element.getparent()
        table_index = list(parent).index(table_element)

        # Insert elements before the table
        for elem_type, level, text, original_content in reversed(elements):
            if elem_type == 'heading':
                # Preserve existing heading style
                new_para = self.doc.add_paragraph(text, style=f'Heading {level}')
                # Move it to the correct position
                new_para_element = new_para._element
                parent.insert(table_index, new_para_element)
            elif elem_type == 'flagged_heading':
                # Create paragraph with original formatting but add a marker
                new_para = self.doc.add_paragraph()

                # Add the marker first
                marker = new_para.add_run("*** STYLE AS HEADING *** ")
                marker.bold = True
                marker.font.color.rgb = RGBColor(255, 0, 0)
                marker.font.size = Pt(10)

                # Add the original text with bold formatting preserved
                content_run = new_para.add_run(text)
                if self.check_paragraph_for_bold(original_content):
                    content_run.bold = True

                # Move it to the correct position
                new_para_element = new_para._element
                parent.insert(table_index, new_para_element)
            elif elem_type == 'title':
                # Preserve existing Title style
                new_para = self.doc.add_paragraph(text, style='Title')
                new_para_element = new_para._element
                parent.insert(table_index, new_para_element)
            elif elem_type == 'flagged_title':
                # Flag text that looks like it should be a Title
                new_para = self.doc.add_paragraph()

                marker = new_para.add_run("*** STYLE AS TITLE *** ")
                marker.bold = True
                marker.font.color.rgb = RGBColor(255, 0, 0)
                marker.font.size = Pt(10)

                content_run = new_para.add_run(text)
                if self.check_paragraph_for_bold(original_content):
                    content_run.bold = True

                new_para_element = new_para._element
                parent.insert(table_index, new_para_element)
            elif elem_type == 'flagged_list':
                # Flag manual list items that need proper list formatting
                new_para = self.doc.add_paragraph()

                marker = new_para.add_run("*** STYLE AS LIST *** ")
                marker.bold = True
                marker.font.color.rgb = RGBColor(255, 0, 0)
                marker.font.size = Pt(10)

                content_run = new_para.add_run(text)

                new_para_element = new_para._element
                parent.insert(table_index, new_para_element)
            elif elem_type == 'flagged_url':
                # Flag raw URLs - insert marker right before the URL
                import re
                new_para = self.doc.add_paragraph()

                # Find URL in text and split around it
                url_pattern = r'(https?://[^\s]+|www\.[^\s]+)'
                parts = re.split(url_pattern, text)

                for i, part in enumerate(parts):
                    if re.match(url_pattern, part):
                        # This is a URL - add marker before it
                        marker = new_para.add_run("*** USE DESCRIPTIVE LINK TEXT *** ")
                        marker.bold = True
                        marker.font.color.rgb = RGBColor(255, 0, 0)
                        marker.font.size = Pt(10)
                        # Add the URL
                        new_para.add_run(part)
                    elif part:
                        # Regular text
                        new_para.add_run(part)

                new_para_element = new_para._element
                parent.insert(table_index, new_para_element)
            elif elem_type == 'flagged_allcaps':
                # Flag ALL CAPS text (accessibility issue)
                new_para = self.doc.add_paragraph()

                marker = new_para.add_run("*** AVOID ALL CAPS *** ")
                marker.bold = True
                marker.font.color.rgb = RGBColor(255, 0, 0)
                marker.font.size = Pt(10)

                content_run = new_para.add_run(text)

                new_para_element = new_para._element
                parent.insert(table_index, new_para_element)
            elif elem_type == 'skipped_heading':
                # Flag heading that skips a level
                new_para = self.doc.add_paragraph()

                marker = new_para.add_run("*** HEADING LEVEL SKIP: Use proper sequence (e.g., H1 → H2, not H1 → H3) *** ")
                marker.bold = True
                marker.font.color.rgb = RGBColor(255, 0, 0)
                marker.font.size = Pt(10)

                # Preserve the heading with its level
                content_run = new_para.add_run(text)
                if self.check_paragraph_for_bold(original_content):
                    content_run.bold = True

                new_para_element = new_para._element
                parent.insert(table_index, new_para_element)
            elif elem_type == 'flagged_link':
                # Flag non-descriptive link text
                new_para = self.doc.add_paragraph()

                marker = new_para.add_run("*** USE DESCRIPTIVE LINK TEXT (not 'click here', 'here', etc.) *** ")
                marker.bold = True
                marker.font.color.rgb = RGBColor(255, 0, 0)
                marker.font.size = Pt(10)

                content_run = new_para.add_run(text)

                new_para_element = new_para._element
                parent.insert(table_index, new_para_element)
            elif elem_type == 'flagged_alignment':
                # Flag manual alignment with tabs/spaces
                new_para = self.doc.add_paragraph()

                marker = new_para.add_run("*** MANUAL ALIGNMENT: Don't use tabs/spaces for alignment. Use tables or proper formatting *** ")
                marker.bold = True
                marker.font.color.rgb = RGBColor(255, 0, 0)
                marker.font.size = Pt(10)

                content_run = new_para.add_run(text)

                new_para_element = new_para._element
                parent.insert(table_index, new_para_element)
            elif elem_type == 'flagged_long':
                # Flag very long paragraph
                new_para = self.doc.add_paragraph()

                marker = new_para.add_run("*** LONG PARAGRAPH: Consider breaking into shorter paragraphs for readability *** ")
                marker.bold = True
                marker.font.color.rgb = RGBColor(255, 0, 0)
                marker.font.size = Pt(10)

                content_run = new_para.add_run(text)

                new_para_element = new_para._element
                parent.insert(table_index, new_para_element)
            elif elem_type == 'table':
                # Preserve nested table - extract and move it
                nested_table_element = original_content._element
                # Remove from current parent (the cell)
                nested_table_parent = nested_table_element.getparent()
                nested_table_parent.remove(nested_table_element)
                # Insert at the new position
                parent.insert(table_index, nested_table_element)
            else:
                # Create new regular paragraph, preserving style
                has_num = self.has_numbering(original_content)

                # Check if original was a list item or has numbering
                if original_content.style.name and 'List' in original_content.style.name:
                    # Preserve list style
                    new_para = self.doc.add_paragraph(text, style=original_content.style.name)
                else:
                    # Regular paragraph (numbering will be applied separately)
                    new_para = self.doc.add_paragraph(text)

                # Preserve Word numbering (bullets/numbers)
                if has_num:
                    self.copy_numbering(original_content, new_para)

                # Preserve alignment and other formatting
                if original_content.alignment:
                    new_para.alignment = original_content.alignment

                # Preserve bold formatting for runs if the whole paragraph was bold
                if self.check_paragraph_for_bold(original_content):
                    for run in new_para.runs:
                        run.bold = True

                # Remove any cell shading (background color) from the paragraph
                # Table cell shading doesn't make sense outside of table context
                if new_para._element.pPr is not None:
                    shd = new_para._element.pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
                    if shd is not None:
                        new_para._element.pPr.remove(shd)

                # Move it to the correct position
                new_para_element = new_para._element
                parent.insert(table_index, new_para_element)

    def has_header_row(self, table) -> bool:
        """Check if table's first row appears to be formatted as headers"""
        if not table.rows:
            return False

        first_row = table.rows[0]

        # Check if first row cells are bold or have different formatting
        bold_count = 0
        total_cells = 0

        for cell in first_row.cells:
            total_cells += 1
            # Check if cell text is bold
            for para in cell.paragraphs:
                for run in para.runs:
                    if run.bold:
                        bold_count += 1
                        break

        # If most cells in first row are bold, it's likely a header row
        if total_cells > 0 and (bold_count / total_cells) > 0.5:
            return True

        return False

    def mark_table_header_issue(self, table) -> bool:
        """
        Add bold red feedback comment to table indicating missing headers.
        Returns True if marker was added.
        """
        if not table.rows or not table.rows[0].cells:
            return False

        # Add feedback marker to first cell
        first_cell = table.rows[0].cells[0]

        # Add feedback as a new paragraph at the beginning
        if first_cell.paragraphs:
            para = first_cell.paragraphs[0]
            # Insert feedback text at the beginning
            run = para.insert_paragraph_before().add_run(
                "*** ACCESSIBILITY ISSUE: This table needs header row (first row should be bold) ***"
            )
            run.bold = True
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.font.size = Pt(10)
            return True

        return False

    def fix_document(self) -> dict:
        """Main method to fix all layout tables in the document"""
        results = {
            'layout_tables_fixed': 0,
            'data_tables_kept': 0,
            'data_tables_marked': 0,
            'total_tables': len(self.doc.tables),
            'fixed_table_details': [],
            'marked_tables': []
        }

        # PHASE 1: Process tables in reverse order (so we can delete them safely)
        print("Phase 1: Converting layout tables...")
        tables = list(self.doc.tables)
        for idx, table in enumerate(reversed(tables)):
            table_num = len(tables) - idx
            is_layout, reason = self.is_layout_table(table)

            if is_layout:
                print(f"  Converting Table {table_num}: {reason}")

                # Convert table to structured content
                elements = self.convert_table_to_structure(table, table_num)

                # Count nested tables and all flagged items
                nested_tables = sum(1 for elem in elements if elem[0] == 'table')
                flagged_headings = sum(1 for elem in elements if elem[0] == 'flagged_heading')
                flagged_titles = sum(1 for elem in elements if elem[0] == 'flagged_title')
                flagged_lists = sum(1 for elem in elements if elem[0] == 'flagged_list')
                flagged_urls = sum(1 for elem in elements if elem[0] == 'flagged_url')
                flagged_allcaps = sum(1 for elem in elements if elem[0] == 'flagged_allcaps')
                skipped_headings = sum(1 for elem in elements if elem[0] == 'skipped_heading')
                flagged_links = sum(1 for elem in elements if elem[0] == 'flagged_link')
                flagged_alignment = sum(1 for elem in elements if elem[0] == 'flagged_alignment')
                flagged_long = sum(1 for elem in elements if elem[0] == 'flagged_long')

                if nested_tables > 0:
                    print(f"    → Preserved {nested_tables} nested table(s)")
                if flagged_headings > 0:
                    print(f"    → Flagged {flagged_headings} unstyled heading(s)")
                if flagged_titles > 0:
                    print(f"    → Flagged {flagged_titles} unstyled title(s)")
                if flagged_lists > 0:
                    print(f"    → Flagged {flagged_lists} manual list item(s)")
                if flagged_urls > 0:
                    print(f"    → Flagged {flagged_urls} raw URL(s)")
                if flagged_allcaps > 0:
                    print(f"    → Flagged {flagged_allcaps} ALL CAPS text(s)")
                if skipped_headings > 0:
                    print(f"    → Flagged {skipped_headings} skipped heading level(s)")
                if flagged_links > 0:
                    print(f"    → Flagged {flagged_links} non-descriptive link(s)")
                if flagged_alignment > 0:
                    print(f"    → Flagged {flagged_alignment} manual alignment(s)")
                if flagged_long > 0:
                    print(f"    → Flagged {flagged_long} long paragraph(s)")

                # Insert new elements before the table
                self.insert_elements_before_table(table, elements)

                # Remove the table
                table_element = table._element
                table_element.getparent().remove(table_element)

                results['layout_tables_fixed'] += 1
                results['headings_flagged'] = results.get('headings_flagged', 0) + flagged_headings
                results['titles_flagged'] = results.get('titles_flagged', 0) + flagged_titles
                results['lists_flagged'] = results.get('lists_flagged', 0) + flagged_lists
                results['urls_flagged'] = results.get('urls_flagged', 0) + flagged_urls
                results['allcaps_flagged'] = results.get('allcaps_flagged', 0) + flagged_allcaps
                results['skipped_headings'] = results.get('skipped_headings', 0) + skipped_headings
                results['links_flagged'] = results.get('links_flagged', 0) + flagged_links
                results['alignment_flagged'] = results.get('alignment_flagged', 0) + flagged_alignment
                results['long_paragraphs'] = results.get('long_paragraphs', 0) + flagged_long
                results['fixed_table_details'].append({
                    'table_num': table_num,
                    'reason': reason,
                    'elements_created': len(elements),
                    'nested_tables_preserved': nested_tables,
                    'headings_flagged': flagged_headings,
                    'titles_flagged': flagged_titles,
                    'lists_flagged': flagged_lists,
                    'urls_flagged': flagged_urls,
                    'allcaps_flagged': flagged_allcaps,
                    'skipped_headings': skipped_headings,
                    'links_flagged': flagged_links,
                    'alignment_flagged': flagged_alignment,
                    'long_paragraphs': flagged_long
                })
            else:
                print(f"  Keeping Table {table_num}: appears to be data table")
                results['data_tables_kept'] += 1

        # PHASE 2: Mark header issues in remaining data tables
        print("")
        print("Phase 2: Marking data table accessibility issues...")
        remaining_tables = list(self.doc.tables)
        for idx, table in enumerate(remaining_tables):
            table_num = idx + 1
            if not self.has_header_row(table):
                if self.mark_table_header_issue(table):
                    print(f"  ✓ Marked Table {table_num}: needs header row")
                    results['data_tables_marked'] += 1
                    results['marked_tables'].append(table_num)

        if results['data_tables_marked'] == 0:
            print("  All data tables have proper headers")
        else:
            print(f"  Marked {results['data_tables_marked']} table(s) with header issues")

        return results

    def save(self, output_path: str):
        """Save the modified document"""
        self.doc.save(output_path)


def main():
    parser = argparse.ArgumentParser(
        description='Convert layout tables to proper headings and paragraphs',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python3 fix_layout_tables.py english215.docx
  python3 fix_layout_tables.py my_syllabus.docx

Output:
  Creates two files:
    - YOUR_SYLLABUS_backup.docx (original file)
    - YOUR_SYLLABUS_fixed.docx (converted file with feedback)
        """
    )

    parser.add_argument('syllabus',
                        help='Path to the syllabus Word document to fix')

    args = parser.parse_args()

    # Validate input file
    input_path = Path(args.syllabus)
    if not input_path.exists():
        print(f"Error: File not found: {input_path}")
        sys.exit(1)

    if not input_path.suffix.lower() in ['.docx', '.doc']:
        print(f"Error: File must be a Word document (.docx or .doc)")
        sys.exit(1)

    # Generate output paths
    stem = input_path.stem
    parent = input_path.parent
    backup_path = parent / f"{stem}_backup.docx"
    fixed_path = parent / f"{stem}_fixed.docx"

    print(f"Processing: {input_path.name}")
    print(f"")

    # Create backup
    print(f"Creating backup: {backup_path.name}")
    shutil.copy2(input_path, backup_path)

    # Load and analyze document
    print(f"Analyzing tables...")
    fixer = LayoutTableFixer(str(input_path))

    # Fix layout tables
    print(f"")
    results = fixer.fix_document()

    # Save fixed document
    print(f"")
    print(f"Saving fixed document: {fixed_path.name}")
    fixer.save(str(fixed_path))

    # Report results
    print(f"")
    print("=" * 60)
    print("RESULTS")
    print("=" * 60)
    print(f"Total tables found: {results['total_tables']}")
    print(f"Layout tables converted: {results['layout_tables_fixed']}")
    print(f"Data tables preserved: {results['data_tables_kept']}")
    print(f"Data table issues marked: {results['data_tables_marked']}")
    print(f"")

    # Accessibility flags summary
    total_flags = (results.get('headings_flagged', 0) + results.get('titles_flagged', 0) +
                   results.get('lists_flagged', 0) + results.get('urls_flagged', 0) +
                   results.get('allcaps_flagged', 0) + results.get('skipped_headings', 0) +
                   results.get('links_flagged', 0) + results.get('alignment_flagged', 0) +
                   results.get('long_paragraphs', 0))
    if total_flags > 0:
        print("Accessibility issues flagged:")
        if results.get('titles_flagged', 0) > 0:
            print(f"  - Unstyled titles: {results.get('titles_flagged', 0)} (*** STYLE AS TITLE ***)")
        if results.get('headings_flagged', 0) > 0:
            print(f"  - Unstyled headings: {results.get('headings_flagged', 0)} (*** STYLE AS HEADING ***)")
        if results.get('skipped_headings', 0) > 0:
            print(f"  - Skipped heading levels: {results.get('skipped_headings', 0)} (*** HEADING LEVEL SKIP ***)")
        if results.get('lists_flagged', 0) > 0:
            print(f"  - Manual lists: {results.get('lists_flagged', 0)} (*** STYLE AS LIST ***)")
        if results.get('urls_flagged', 0) > 0:
            print(f"  - Raw URLs: {results.get('urls_flagged', 0)} (*** USE DESCRIPTIVE LINK TEXT ***)")
        if results.get('links_flagged', 0) > 0:
            print(f"  - Non-descriptive links: {results.get('links_flagged', 0)} (*** USE DESCRIPTIVE LINK TEXT ***)")
        if results.get('alignment_flagged', 0) > 0:
            print(f"  - Manual alignment: {results.get('alignment_flagged', 0)} (*** MANUAL ALIGNMENT ***)")
        if results.get('allcaps_flagged', 0) > 0:
            print(f"  - ALL CAPS text: {results.get('allcaps_flagged', 0)} (*** AVOID ALL CAPS ***)")
        if results.get('long_paragraphs', 0) > 0:
            print(f"  - Long paragraphs: {results.get('long_paragraphs', 0)} (*** LONG PARAGRAPH ***)")
        print(f"")

    if results['fixed_table_details']:
        print("Converted tables:")
        for detail in results['fixed_table_details']:
            print(f"  - Table {detail['table_num']}: Created {detail['elements_created']} elements")
            if detail.get('nested_tables_preserved', 0) > 0:
                print(f"    Preserved {detail['nested_tables_preserved']} nested table(s)")
            # Show flagged items for this table
            flags_in_table = []
            if detail.get('titles_flagged', 0) > 0:
                flags_in_table.append(f"{detail['titles_flagged']} title(s)")
            if detail.get('headings_flagged', 0) > 0:
                flags_in_table.append(f"{detail['headings_flagged']} heading(s)")
            if detail.get('skipped_headings', 0) > 0:
                flags_in_table.append(f"{detail['skipped_headings']} skipped heading(s)")
            if detail.get('lists_flagged', 0) > 0:
                flags_in_table.append(f"{detail['lists_flagged']} list(s)")
            if detail.get('urls_flagged', 0) > 0:
                flags_in_table.append(f"{detail['urls_flagged']} URL(s)")
            if detail.get('links_flagged', 0) > 0:
                flags_in_table.append(f"{detail['links_flagged']} non-descriptive link(s)")
            if detail.get('alignment_flagged', 0) > 0:
                flags_in_table.append(f"{detail['alignment_flagged']} manual alignment(s)")
            if detail.get('allcaps_flagged', 0) > 0:
                flags_in_table.append(f"{detail['allcaps_flagged']} ALL CAPS")
            if detail.get('long_paragraphs', 0) > 0:
                flags_in_table.append(f"{detail['long_paragraphs']} long paragraph(s)")
            if flags_in_table:
                print(f"    Flagged: {', '.join(flags_in_table)}")
            print(f"    Reason: {detail['reason']}")
        print(f"")

    if results['marked_tables']:
        print("Data tables marked with accessibility issues:")
        for table_num in results['marked_tables']:
            print(f"  - Table {table_num}: Missing header row (feedback added)")

    print(f"")
    print("=" * 60)
    print("NEXT STEPS")
    print("=" * 60)
    print(f"1. Open the fixed document: {fixed_path.name}")
    print(f"2. Look for bold red text marking issues")
    print(f"3. Fix the marked issues")
    print(f"4. Run syllabus checker to verify:")
    print(f"   python3 syllabus_checker.py {fixed_path.name}")
    print(f"5. Original backup saved at: {backup_path.name}")
    print("=" * 60)


if __name__ == "__main__":
    main()
