#!/usr/bin/env python3
"""
Test script to verify grading scale table is not flagged as a layout table
"""

from docx import Document
from docx.shared import Pt, RGBColor
from syllabus_checker import SyllabusChecker

# Create a test document with a grading scale table
doc = Document()

# Add a title
doc.add_heading('Test Syllabus', 1)

# Create the grading scale table (2 columns, 13 rows + header)
table = doc.add_table(rows=14, cols=2)
table.style = 'Light Grid Accent 1'

# Add headers (make them bold)
header_cells = table.rows[0].cells
header_cells[0].text = 'Grade'
header_cells[1].text = 'Scale'

# Make headers bold
for cell in header_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

# Add grading scale data
grades_data = [
    ('A', '93-100'),
    ('A-', '90-92'),
    ('B+', '87-89'),
    ('B', '83-86'),
    ('B-', '80-82'),
    ('C+', '77-79'),
    ('C', '73-76'),
    ('C-', '70-72'),
    ('D+', '67-69'),
    ('D', '63-66'),
    ('D-', '60-62'),
    ('F', '0-59')
]

for i, (grade, scale) in enumerate(grades_data, start=1):
    row_cells = table.rows[i].cells
    row_cells[0].text = grade
    row_cells[1].text = scale

# Save test document
test_doc_path = 'test_grading_table.docx'
doc.save(test_doc_path)
print(f"Created test document: {test_doc_path}")

# Now test it with the checker
print("\nRunning syllabus checker on test document...\n")

# We need a template - let's use the existing one
template_path = 'Uniform-Syllabus-Template-1.docx'

checker = SyllabusChecker(template_path, test_doc_path)
checker.run_all_checks()

# Check for layout table issues
layout_issues = [issue for issue in checker.issues if issue.issue_type == "LAYOUT_TABLE"]

print("=" * 60)
print("TEST RESULTS")
print("=" * 60)
print(f"Total tables in document: {len(checker.target_doc.tables)}")
print(f"Layout table issues found: {len(layout_issues)}")
print()

if layout_issues:
    print("❌ FAILED - Grading table was incorrectly flagged as layout table:")
    for issue in layout_issues:
        print(f"  - {issue.description}")
else:
    print("✅ PASSED - Grading table correctly recognized as data table!")
    print("   (Tables with proper headers are no longer flagged)")

print("=" * 60)
