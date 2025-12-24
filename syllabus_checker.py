#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Syllabus Checker - Assesses syllabi against UWM template standards
Checks for missing sections and accessibility issues
Analyzes content both in document body AND within table cells
Generates both a text report and a marked-up Word document with tracked issues
"""

import argparse
import re
from typing import List, Dict, Tuple, Optional
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def calculate_relative_luminance(rgb: Tuple[int, int, int]) -> float:
    """Calculate relative luminance for a color (WCAG formula)"""
    def adjust_channel(channel: int) -> float:
        c = channel / 255.0
        if c <= 0.03928:
            return c / 12.92
        else:
            return ((c + 0.055) / 1.055) ** 2.4

    r, g, b = rgb
    r_adj = adjust_channel(r)
    g_adj = adjust_channel(g)
    b_adj = adjust_channel(b)

    return 0.2126 * r_adj + 0.7152 * g_adj + 0.0722 * b_adj


def calculate_contrast_ratio(color1: Tuple[int, int, int], color2: Tuple[int, int, int]) -> float:
    """Calculate contrast ratio between two colors (WCAG formula)"""
    lum1 = calculate_relative_luminance(color1)
    lum2 = calculate_relative_luminance(color2)

    lighter = max(lum1, lum2)
    darker = min(lum1, lum2)

    return (lighter + 0.05) / (darker + 0.05)


class ParagraphInfo:
    """Stores information about a paragraph and its location"""
    def __init__(self, paragraph, index: int, location: str,
                 table_idx: Optional[int] = None,
                 row_idx: Optional[int] = None,
                 cell_idx: Optional[int] = None):
        self.paragraph = paragraph
        self.index = index  # Global paragraph index
        self.location = location  # Human-readable location
        self.table_idx = table_idx
        self.row_idx = row_idx
        self.cell_idx = cell_idx


class AccessibilityIssue:
    """Represents a specific accessibility issue in the document"""
    def __init__(self, issue_type: str, description: str, location: str,
                 para_info: Optional[ParagraphInfo] = None,
                 table_index: Optional[int] = None):
        self.issue_type = issue_type
        self.description = description
        self.location = location
        self.para_info = para_info
        self.table_index = table_index


class SyllabusChecker:
    """Analyzes a syllabus document for completeness and accessibility"""

    # Required sections based on UWM template
    REQUIRED_SECTIONS = {
        "Course Title": ["course title", "course name"],
        "Instructor": ["instructor", "professor", "teacher"],
        "Welcome Statement": ["welcome"],
        "Course Overview": ["course overview", "overview", "course description"],
        "Course Objectives": ["course objectives", "objectives", "goals"],
        "Student Learning Outcomes": ["student learning outcomes", "learning outcomes", "outcomes"],
        "Assessment": ["assessment"],
        "Prerequisites": ["prerequisites", "prerequisite", "required skills"],
        "Course Modality": ["course modality", "modality", "format", "meeting times"],
        "Course Materials": ["course materials", "materials", "textbook", "required materials"],
        "Time Investment": ["time investment", "expected workload", "time commitment"],
        "Assignments & Grading": ["assignments", "grading", "grade"],
        "Grading Scheme": ["grading scheme", "grade breakdown"],
        "Grading Scale": ["grading scale", "grade scale"],
        "Course Policies": ["course policies", "policies"],
        "University Policies": ["university policies", "uwm policies"],
        "Resources": ["resources", "student resources"],
        "Academic Supports": ["academic supports", "tutoring", "academic resources"],
        "Important Dates": ["important dates", "key dates", "deadlines"],
        "Calendar": ["calendar", "schedule"],
    }

    def __init__(self, template_path: str, target_path: str):
        """Initialize checker with template and target syllabus paths"""
        self.template_doc = Document(template_path)
        self.target_doc = Document(target_path)
        self.target_path = target_path
        self.template_sections = self._extract_sections(self.template_doc)
        self.target_sections = self._extract_sections(self.target_doc)
        self.all_paragraphs = self._get_all_paragraphs()  # NEW: Get all paragraphs including from tables
        self.issues = []  # Store all detected issues

    def _get_all_paragraphs(self) -> List[ParagraphInfo]:
        """Extract ALL paragraphs from document, including those inside table cells"""
        all_paras = []
        global_index = 0

        # Get top-level paragraphs
        for para in self.target_doc.paragraphs:
            para_info = ParagraphInfo(
                paragraph=para,
                index=global_index,
                location=f"Paragraph {global_index + 1}"
            )
            all_paras.append(para_info)
            global_index += 1

        # Get paragraphs from within tables
        for table_idx, table in enumerate(self.target_doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para in cell.paragraphs:
                        para_info = ParagraphInfo(
                            paragraph=para,
                            index=global_index,
                            location=f"Table {table_idx + 1}, Row {row_idx + 1}, Cell {cell_idx + 1}",
                            table_idx=table_idx,
                            row_idx=row_idx,
                            cell_idx=cell_idx
                        )
                        all_paras.append(para_info)
                        global_index += 1

        return all_paras

    def _extract_sections(self, doc: Document) -> List[Dict]:
        """Extract headings and their content from document (including tables)"""
        sections = []

        # Check top-level paragraphs
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                level = int(para.style.name.split()[-1])
                sections.append({
                    'text': para.text.strip(),
                    'level': level,
                    'style': para.style.name,
                    'location': 'top-level'
                })

        # Check paragraphs in tables
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para in cell.paragraphs:
                        if para.style.name.startswith('Heading'):
                            level = int(para.style.name.split()[-1])
                            sections.append({
                                'text': para.text.strip(),
                                'level': level,
                                'style': para.style.name,
                                'location': f'Table {table_idx + 1}, Row {row_idx + 1}, Cell {cell_idx + 1}'
                            })

        return sections

    def _find_section(self, section_name: str, search_terms: List[str],
                     sections: List[Dict]) -> bool:
        """Check if a section exists in the document"""
        # Check in extracted heading sections
        for section in sections:
            section_text = section['text'].lower()
            if any(term in section_text for term in search_terms):
                return True

        # Also check in ALL paragraph text (using our comprehensive list)
        for para_info in self.all_paragraphs:
            para_text = para_info.paragraph.text.lower()
            if any(term in para_text for term in search_terms):
                # Only count if it's at the start of a line/paragraph (likely a heading)
                if len(para_info.paragraph.text.strip()) < 100:  # Likely a heading, not body text
                    return True
        return False

    def check_missing_sections(self) -> List[str]:
        """Identify missing required sections"""
        missing = []
        for section, search_terms in self.REQUIRED_SECTIONS.items():
            if not self._find_section(section, search_terms, self.target_sections):
                missing.append(section)
        return missing

    def check_unstyled_headings(self) -> List[AccessibilityIssue]:
        """Detect text that looks like headings but doesn't use heading styles"""
        unstyled_headings = []

        # Track last heading level to suggest appropriate level
        last_heading_level = 0
        last_heading_location = None

        for para_info in self.all_paragraphs:
            para = para_info.paragraph

            # Track actual heading styles we encounter
            if para.style.name.startswith('Heading'):
                last_heading_level = int(para.style.name.split()[-1])
                last_heading_location = para_info.location
                continue

            text = para.text.strip()

            # Skip empty paragraphs or very long paragraphs
            if not text or len(text) > 150:
                continue

            # FILTER OUT FALSE POSITIVES

            # Skip list items (likely just emphasized list text, not headings)
            if para.style.name and 'List' in para.style.name:
                continue

            # Skip if starts with bullet/number (manual list items)
            if text and text[0] in ['•', '·', '-', '*', '○', '■']:
                continue

            # Skip if looks like a numbered item
            if len(text) > 0 and text[0].isdigit() and '.' in text[:5]:
                continue

            # Skip very short text (1-2 chars) - likely labels, not headings
            if len(text) <= 2:
                continue

            # Skip if only PART of the text is bold (inline emphasis, not heading)
            if para.runs:
                bold_chars = sum(len(run.text) for run in para.runs if run.bold)
                total_chars = len(text)
                bold_ratio = bold_chars / total_chars if total_chars > 0 else 0

                # If less than 80% bold, it's probably inline emphasis
                if 0 < bold_ratio < 0.8:
                    continue

            # Skip if ends with certain punctuation (likely emphasized text, not heading)
            if text.endswith(('.', '!', ',', ';', ':')):
                continue

            # Skip text that starts with deadline/assignment indicators (not headings)
            # Examples: "Due by January 15", "Submit your assignment", "Complete the reading"
            text_lower = text.lower()
            deadline_prefixes = [
                'due by', 'due on', 'due date', 'due:', 'submit', 'turn in',
                'complete', 'finish', 'read', 'prepare', 'bring to class'
            ]
            if any(text_lower.startswith(prefix) for prefix in deadline_prefixes):
                continue

            # Skip "Label: Description" patterns (activity labels, not headings)
            # Examples: "In-class writing: Business packet", "Reading: Chapter 5", "Homework: Read chapter 5"
            if ':' in text:
                # Check if this is a label:description pattern (not a section heading)
                parts = text.split(':', 1)
                if len(parts) == 2:
                    label = parts[0].strip()
                    description = parts[1].strip()

                    # Common activity/assignment labels (not major section headings)
                    activity_labels = [
                        'in-class', 'homework', 'assignment', 'due', 'reading', 'activity',
                        'discussion', 'group work', 'lab', 'quiz', 'exam', 'project',
                        'presentation', 'review', 'workshop', 'lecture', 'tutorial'
                    ]

                    # If the label is short and matches activity patterns, skip it
                    label_lower = label.lower()
                    if len(label) < 30 and description:  # Has both label and description
                        # Check if it's an activity label
                        if any(activity_word in label_lower for activity_word in activity_labels):
                            continue

                        # Also skip if the label is short and description contains specific items
                        # (e.g., "Reading: Chapter 5", "Due: Assignment 3")
                        if len(label.split()) <= 4 and len(description.split()) <= 10:
                            # This is likely a label:description, not a major section heading
                            continue

            # NOW CHECK FOR HEADING CHARACTERISTICS
            is_likely_heading = False
            reasons = []

            # Check for bold text (full paragraph bold)
            if para.runs:
                bold_chars = sum(len(run.text) for run in para.runs if run.bold)
                total_chars = len(text)
                if total_chars > 0 and bold_chars / total_chars >= 0.8:
                    # Additional check: is it standalone and short?
                    words = text.split()
                    if len(words) <= 10:  # Headings are typically short
                        is_likely_heading = True
                        reasons.append("bold text")

            # Check for all caps (but not single words like "I" or "A")
            if text.isupper() and len(text) > 3:
                is_likely_heading = True
                reasons.append("all caps")

            # Check for title case + no ending punctuation + short
            if len(text) < 100:
                words = text.split()
                if 2 <= len(words) <= 8:  # Multi-word but not too long
                    # Check for larger font size
                    has_large_font = False
                    for run in para.runs:
                        if run.font.size and run.font.size.pt > 12:
                            has_large_font = True
                            reasons.append("large font")
                            break

                    if has_large_font or (bold_chars > 0 and len(words) <= 6):
                        is_likely_heading = True

            if is_likely_heading and reasons:
                # Determine recommended heading level based on context
                recommended_level = self._recommend_heading_level(
                    para_info, text, last_heading_level, last_heading_location
                )

                issue = AccessibilityIssue(
                    issue_type=f"SHOULD BE HEADING {recommended_level}",
                    description=f"Text appears to be a heading ({', '.join(reasons)}): \"{text[:60]}...\"",
                    location=para_info.location,
                    para_info=para_info
                )
                unstyled_headings.append(issue)

        return unstyled_headings

    def _recommend_heading_level(self, para_info: ParagraphInfo, text: str,
                                  last_heading_level: int, last_heading_location: str) -> int:
        """Recommend appropriate heading level based on context"""

        # Check if text looks like a major section title
        major_section_keywords = [
            'overview', 'introduction', 'objectives', 'materials', 'assignments',
            'grading', 'policies', 'schedule', 'calendar', 'resources'
        ]

        text_lower = text.lower()
        is_major_section = any(keyword in text_lower for keyword in major_section_keywords)

        # If in a table (likely nested content)
        if para_info.table_idx is not None:
            # If it's a recognized major section, could be H1
            if is_major_section and len(text.split()) <= 3:
                return 1
            # Otherwise likely a subsection
            return 2

        # Top-level paragraph
        else:
            # First heading in document or major section
            if last_heading_level == 0 or is_major_section:
                return 1
            # Following an H1, likely also H1 or H2
            elif last_heading_level == 1:
                # If it looks like a subsection, suggest H2
                if len(text.split()) > 3:
                    return 2
                return 1
            # Following an H2 or deeper
            else:
                return last_heading_level

        return 1  # Default to H1 if unsure

    def _table_has_header_row(self, table) -> bool:
        """Helper: Check if a table has a proper header row"""
        if not table.rows:
            return False

        first_row = table.rows[0]
        num_cells = len(first_row.cells)

        # Check multiple indicators of header rows
        bold_count = 0
        shaded_count = 0
        large_font_count = 0

        for cell in first_row.cells:
            # Check for bold text
            has_bold = False
            has_large_font = False

            for para in cell.paragraphs:
                for run in para.runs:
                    if run.bold:
                        has_bold = True
                    # Check for larger font (>11pt suggests header)
                    if run.font.size and run.font.size.pt > 11:
                        has_large_font = True

            if has_bold:
                bold_count += 1
            if has_large_font:
                large_font_count += 1

            # Check for cell shading (background color)
            # Access the cell's XML to check for shading
            try:
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                shd = tcPr.find(qn('w:shd'), tcPr.nsmap)
                if shd is not None:
                    fill = shd.get(qn('w:fill'))
                    # If there's a fill color (not 'auto'), likely a header
                    if fill and fill != 'auto' and fill != 'FFFFFF':
                        shaded_count += 1
            except:
                pass  # If we can't check shading, skip it

        # Header row if most cells have ANY of these indicators:
        # - Bold text (50%+)
        # - Shading (50%+)
        # - Large font (50%+)
        if bold_count >= num_cells * 0.5:
            return True
        if shaded_count >= num_cells * 0.5:
            return True
        if large_font_count >= num_cells * 0.5:
            return True

        return False

    def check_table_headers(self) -> List[AccessibilityIssue]:
        """Check if tables have proper header rows"""
        table_issues = []

        for i, table in enumerate(self.target_doc.tables):
            has_header = self._table_has_header_row(table)

            if not has_header and len(table.rows) > 1:
                # Get preview of table content
                preview = " | ".join([cell.text.strip()[:20] for cell in table.rows[0].cells])

                issue = AccessibilityIssue(
                    issue_type="TABLE_NO_HEADER",
                    description=f"Table {i+1} appears to lack header row. First row: \"{preview}...\"",
                    location=f"Table {i+1}",
                    table_index=i
                )
                table_issues.append(issue)

        return table_issues

    def check_layout_tables(self) -> List[AccessibilityIssue]:
        """Identify tables used for layout rather than tabular data"""
        layout_tables = []

        for i, table in enumerate(self.target_doc.tables):
            is_layout_table = False
            reasons = []

            rows = len(table.rows)
            cols = len(table.columns)

            # Check if table has headers (data tables with headers are generally OK)
            has_header = self._table_has_header_row(table)

            # Check for single column/row (almost always layout)
            if cols == 1 or rows == 1:
                is_layout_table = True
                reasons.append("single column/row")

            # Check for cells with excessive text
            cell_lengths = []
            empty_cells = 0
            very_long_cells = 0

            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    cell_len = len(cell_text)
                    cell_lengths.append(cell_len)

                    if cell_len == 0:
                        empty_cells += 1
                    elif cell_len > 500:
                        very_long_cells += 1

            total_cells = len(cell_lengths)

            # Only flag empty cells if table lacks headers (data tables can have sparse data)
            if not has_header and total_cells > 0 and empty_cells / total_cells > 0.3:
                is_layout_table = True
                reasons.append(f"{empty_cells}/{total_cells} empty cells")

            # If cells have very long text, likely layout (not data)
            # But only if no headers (tables with headers might legitimately have long cells)
            if not has_header and very_long_cells > 0:
                is_layout_table = True
                reasons.append(f"{very_long_cells} cells with >500 chars")

            # Check for inconsistent cell lengths only if no headers
            # (data tables can have varying cell content)
            if not has_header and len(cell_lengths) > 4:
                import statistics
                try:
                    stdev = statistics.stdev(cell_lengths)
                    mean = statistics.mean(cell_lengths)
                    if mean > 0 and stdev / mean > 2:  # High variance
                        is_layout_table = True
                        reasons.append("inconsistent cell content lengths")
                except:
                    pass

            # Check for two-column "label: value" pattern
            # BUT: Don't flag if the table has proper headers (first row bold)
            if cols == 2 and rows > 2:
                # Use helper to check if table has header row
                has_header_row = self._table_has_header_row(table)

                # Only check for label-value pattern if NO header row
                if not has_header_row:
                    label_value_pattern = 0
                    for row in table.rows:
                        if len(row.cells) >= 2:
                            cell1_len = len(row.cells[0].text.strip())
                            cell2_len = len(row.cells[1].text.strip())
                            if 0 < cell1_len < 50 and cell2_len > cell1_len:
                                label_value_pattern += 1

                    if label_value_pattern / rows > 0.6:
                        is_layout_table = True
                        reasons.append("label-value pattern detected")

            if is_layout_table:
                issue = AccessibilityIssue(
                    issue_type="LAYOUT_TABLE",
                    description=f"Table {i+1} appears to be used for layout ({'; '.join(reasons)}). "
                                f"Size: {rows} rows × {cols} cols. Use headings and paragraphs instead.",
                    location=f"Table {i+1}",
                    table_index=i
                )
                layout_tables.append(issue)

        return layout_tables

    def check_heading_hierarchy(self) -> List[AccessibilityIssue]:
        """Check for improper heading hierarchy (e.g., too many H1s, H1s under H1s)"""
        hierarchy_issues = []

        # Track heading sequence to detect hierarchy problems
        last_heading_level = 0
        h1_count = 0

        for para_info in self.all_paragraphs:
            para = para_info.paragraph
            if para.style.name.startswith('Heading'):
                level = int(para.style.name.split()[-1])
                text = para.text.strip()

                # Count H1s
                if level == 1:
                    h1_count += 1

                    # If we've seen an H1 before and now see another H1, check context
                    # In table cells, multiple H1s might indicate subsections that should be H2
                    if h1_count > 3 and para_info.table_idx is not None:
                        # This H1 is inside a table and there are many H1s
                        # It's likely a subsection that should be H2
                        issue = AccessibilityIssue(
                            issue_type="SHOULD BE HEADING 2",
                            description=f"'{text}' uses Heading 1 but appears to be a subsection (should be Heading 2)",
                            location=para_info.location,
                            para_info=para_info
                        )
                        hierarchy_issues.append(issue)

                # Check for skipped levels (e.g., H1 -> H3, skipping H2)
                if last_heading_level > 0 and level > last_heading_level + 1:
                    # It should be one level below the previous heading
                    recommended_level = last_heading_level + 1
                    issue = AccessibilityIssue(
                        issue_type=f"SHOULD BE HEADING {recommended_level}",
                        description=f"'{text}' jumps from Heading {last_heading_level} to Heading {level} (should be Heading {recommended_level})",
                        location=para_info.location,
                        para_info=para_info
                    )
                    hierarchy_issues.append(issue)

                last_heading_level = level

        return hierarchy_issues

    def check_heading_structure(self) -> Dict[str, any]:
        """Analyze heading usage and structure"""
        issues = []
        heading_counts = {}

        # Count headings by level (now checking all paragraphs)
        for para_info in self.all_paragraphs:
            para = para_info.paragraph
            if para.style.name.startswith('Heading'):
                level = para.style.name
                heading_counts[level] = heading_counts.get(level, 0) + 1

        # Check for proper hierarchy
        if not heading_counts:
            issues.append("CRITICAL: No heading styles used. Document should use Heading 1, Heading 2, etc.")
        elif len(heading_counts) < 2:
            issues.append("WARNING: Limited heading hierarchy. Use multiple heading levels for better structure.")

        # Check if Heading 1 exists
        if 'Heading 1' not in heading_counts:
            issues.append("WARNING: No Heading 1 found. Major sections should use Heading 1.")

        # Check for proper nesting
        levels_used = sorted([int(h.split()[-1]) for h in heading_counts.keys()])
        if levels_used and levels_used[0] != 1:
            issues.append("WARNING: Document should start with Heading 1, not Heading {}.".format(levels_used[0]))

        # Warn about excessive H1s only if there's evidence of hierarchy problems
        if 'Heading 1' in heading_counts:
            h1_count = heading_counts['Heading 1']
            h2_count = heading_counts.get('Heading 2', 0)

            # Only warn if:
            # 1. Many H1s (>10) with very few H2s (suggests bad hierarchy)
            # 2. Or many H1s (>15) total (likely overuse regardless)
            if h1_count > 10 and h2_count < 3:
                issues.append(
                    f"WARNING: Document has {h1_count} Heading 1s but only {h2_count} Heading 2s. "
                    "Many H1s should likely be H2 (subsections)."
                )
            elif h1_count > 15:
                issues.append(
                    f"WARNING: Document has {h1_count} Heading 1s. This is excessive - "
                    "consider using Heading 2 for subsections."
                )

        return {
            'counts': heading_counts,
            'issues': issues
        }

    def check_table_usage(self) -> Dict[str, any]:
        """Analyze table usage for accessibility"""
        issues = []
        tables = self.target_doc.tables

        if len(tables) == 0:
            return {'count': 0, 'issues': []}

        complex_tables = []
        for i, table in enumerate(tables):
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if len(cell_text) > 500:
                        complex_tables.append(i)
                        issues.append(
                            f"WARNING: Table {i+1} contains cells with excessive text ({len(cell_text)} chars). "
                            "Consider using regular paragraphs with headings instead."
                        )
                        break

        if len(tables) > 5:  # Increased threshold from 2 to 5
            issues.append(
                f"WARNING: Document contains {len(tables)} tables. "
                "Excessive table usage can harm accessibility. Consider using headings and lists."
            )

        # Check for large tables, but only warn if they don't have headers
        # (tables with headers are likely legitimate data tables)
        for i, table in enumerate(tables):
            if len(table.rows) > 10 or len(table.columns) > 5:
                # Check if it has a header row (data tables with headers are OK)
                has_header = self._table_has_header_row(table)
                if not has_header:
                    issues.append(
                        f"WARNING: Table {i+1} is large ({len(table.rows)} rows, {len(table.columns)} cols) "
                        "and lacks proper header row. This may indicate using tables for layout."
                    )

        return {
            'count': len(tables),
            'complex_tables': complex_tables,
            'issues': issues
        }

    def check_hyperlinks(self) -> Dict[str, any]:
        """Check for accessible hyperlink practices"""
        issues = []
        hyperlinks = []

        # Check all paragraphs (including in tables)
        for para_info in self.all_paragraphs:
            para = para_info.paragraph
            for run in para.runs:
                if run.font.underline or run.font.color.rgb:
                    if re.search(r'https?://', run.text):
                        hyperlinks.append(run.text)
                        if run.text.startswith('http'):
                            issues.append(
                                f"WARNING: Hyperlink displays URL directly: '{run.text[:50]}...'. "
                                "Use descriptive link text instead."
                            )

        return {
            'count': len(hyperlinks),
            'issues': issues
        }

    def check_images(self) -> Dict[str, any]:
        """Check for image alt text"""
        issues = []
        images_without_alt = []

        for rel in self.target_doc.part.rels.values():
            if "image" in rel.target_ref:
                images_without_alt.append(rel.target_ref)

        if images_without_alt:
            issues.append(
                f"INFO: Document contains {len(images_without_alt)} image(s). "
                "Verify all images have descriptive alt text."
            )

        return {
            'count': len(images_without_alt),
            'issues': issues
        }

    def check_text_formatting(self) -> Dict[str, any]:
        """Check for proper use of styles vs. manual formatting"""
        issues = []
        manual_formatting_count = 0

        for para_info in self.all_paragraphs:
            para = para_info.paragraph
            text = para.text.strip()

            # Only check Normal styled paragraphs
            if para.style.name == 'Normal' and para.runs and text:
                # Skip very short text (likely labels, not headings)
                if len(text) <= 3:
                    continue

                # Skip text that ends with sentence punctuation (likely inline emphasis)
                if text.endswith(('.', '!', '?', ',', ';', ':')):
                    continue

                # Check if MOST of the paragraph is bold (>80% = likely a heading, not inline emphasis)
                bold_chars = sum(len(run.text) for run in para.runs if run.bold)
                total_chars = len(text)

                if total_chars > 0:
                    bold_ratio = bold_chars / total_chars

                    # Only flag if most of the text is bold AND it's relatively short (heading-like)
                    if bold_ratio > 0.8 and len(text) < 100:
                        # Additional check: skip if it's in the middle of a sentence
                        # (look for lowercase start or mid-sentence indicators)
                        if text and text[0].isupper():
                            manual_formatting_count += 1

        if manual_formatting_count > 5:
            issues.append(
                f"WARNING: Found {manual_formatting_count} instances of manual bold formatting that look like headings. "
                "Use Heading styles instead of bold text for section headers."
            )

        return {
            'manual_formatting_count': manual_formatting_count,
            'issues': issues
        }

    def check_list_usage(self) -> Dict[str, any]:
        """Check for proper list formatting"""
        issues = []
        list_count = 0
        manual_list_count = 0

        for para_info in self.all_paragraphs:
            para = para_info.paragraph
            if 'List' in para.style.name:
                list_count += 1
            elif para.style.name == 'Normal':
                text = para.text.strip()
                if text and text[0] in ['-', '*', '•', '·']:
                    manual_list_count += 1

        if manual_list_count > 3:
            issues.append(
                f"WARNING: Found {manual_list_count} manually formatted list items. "
                "Use built-in list styles for better accessibility."
            )

        return {
            'list_count': list_count,
            'manual_list_count': manual_list_count,
            'issues': issues
        }

    def check_manual_alignment(self) -> List[AccessibilityIssue]:
        """Check for manual alignment using tabs and multiple spaces (pseudo-tables)"""
        issues = []

        # Get list of paragraphs already flagged as heading issues to avoid double-flagging
        heading_issue_paras = set()
        for issue in self.issues:
            if issue.issue_type.startswith("SHOULD BE HEADING") and issue.para_info:
                heading_issue_paras.add(id(issue.para_info.paragraph))

        for para_info in self.all_paragraphs:
            para = para_info.paragraph
            text = para.text

            # Skip if already flagged as a heading issue
            if id(para) in heading_issue_paras:
                continue

            # Skip very short text (unlikely to be pseudo-tables)
            if len(text.strip()) < 10:
                continue

            is_pseudo_table = False
            reasons = []

            # Check for tabs used for columnar alignment
            # Only flag if there's substantive content (5+ chars) on both sides of a tab
            if '\t' in text:
                tab_parts = text.split('\t')
                for i in range(len(tab_parts) - 1):
                    left_side = tab_parts[i].strip()
                    right_side = tab_parts[i + 1].strip()
                    # If both sides have content, it's likely a pseudo-table
                    if len(left_side) >= 5 and len(right_side) >= 3:
                        is_pseudo_table = True
                        reasons.append("tabs for columnar alignment")
                        break

            # Check for multiple spaces used for alignment
            # Only flag if: 4+ consecutive spaces OR multiple instances of 2+ spaces
            space_matches = list(re.finditer(r'  +', text))
            if space_matches:
                # Flag if any match has 4+ spaces (clearly intentional alignment, not just sentence spacing)
                # But ignore trailing spaces at the end
                significant_matches = [m for m in space_matches if m.end() < len(text.rstrip())]

                if any(len(match.group()) >= 4 for match in significant_matches):
                    is_pseudo_table = True
                    reasons.append("multiple spaces for alignment")
                # Or if there are multiple occurrences of 2+ spaces (suggesting columns)
                elif len(significant_matches) >= 2:
                    # Additional check: make sure they're not all just sentence spacing
                    # If all matches are exactly after periods, it's likely just double-spacing between sentences
                    non_sentence_spacing = []
                    for match in significant_matches:
                        # Check if this is sentence spacing (period before the spaces)
                        before_pos = match.start() - 1
                        if before_pos >= 0 and text[before_pos] not in '.!?':
                            non_sentence_spacing.append(match)

                    if len(non_sentence_spacing) >= 2:
                        is_pseudo_table = True
                        reasons.append("multiple spacing gaps (columnar layout)")

            if is_pseudo_table:
                # Create a preview of the problematic text
                preview = text[:60] + "..." if len(text) > 60 else text

                issue = AccessibilityIssue(
                    issue_type="PSEUDO_TABLE",
                    description=f"Uses {' and '.join(reasons)}: \"{preview}\"",
                    location=para_info.location,
                    para_info=para_info
                )
                issues.append(issue)

        return issues

    def check_font_sizes(self) -> List[AccessibilityIssue]:
        """Check for font sizes below 11pt (inaccessible for body text)"""
        issues = []

        for para_info in self.all_paragraphs:
            para = para_info.paragraph
            text = para.text.strip()

            # Skip empty paragraphs
            if not text:
                continue

            # Check each run for small font sizes
            for run in para.runs:
                if run.font.size and run.font.size.pt < 11:
                    # Get a preview of the text with small font
                    run_text = run.text.strip()
                    if run_text:  # Only flag if there's actual text content
                        preview = text[:60] + "..." if len(text) > 60 else text

                        issue = AccessibilityIssue(
                            issue_type="SMALL_FONT",
                            description=f"Font size {run.font.size.pt}pt (minimum 11pt): \"{preview}\"",
                            location=para_info.location,
                            para_info=para_info
                        )
                        issues.append(issue)
                        break  # Only flag paragraph once

        return issues

    def check_decorative_fonts(self) -> List[AccessibilityIssue]:
        """Detect decorative or inaccessible fonts that are hard to read"""
        issues = []

        # List of decorative/display fonts that are hard to read or inaccessible
        # These include script fonts, decorative fonts, and fonts known for poor readability
        decorative_fonts = {
            # Script/Handwriting fonts
            'brush script', 'lucida handwriting', 'mistral', 'monotype corsiva',
            'palace script', 'vivaldi', 'edwardian script', 'freestyle script',
            'french script', 'kunstler script', 'script', 'cursive',

            # Display/Decorative fonts
            'curlz', 'jokerman', 'ravie', 'showcard gothic', 'snap itc',
            'stencil', 'algerian', 'broadway', 'old english', 'chiller',
            'harrington', 'papyrus', 'comic sans', 'impact',

            # Gothic/Blackletter fonts
            'blackadder', 'fraktur', 'old english text', 'blackletter',

            # Comic/Novelty fonts
            'comic sans ms', 'comic sans', 'kristen itc', 'tempus sans itc',
            'gigi', 'marker', 'jester'
        }

        # Track fonts we've already flagged to avoid duplicate warnings
        flagged_fonts = set()

        for para_info in self.all_paragraphs:
            para = para_info.paragraph
            text = para.text.strip()

            # Skip empty paragraphs
            if not text:
                continue

            # Check each run for decorative fonts
            for run in para.runs:
                run_text = run.text.strip()
                if not run_text:
                    continue

                font_name = run.font.name
                if font_name:
                    font_name_lower = font_name.lower()

                    # Check if this font is in our decorative list
                    is_decorative = False
                    matched_font = None
                    for dec_font in decorative_fonts:
                        if dec_font in font_name_lower:
                            is_decorative = True
                            matched_font = dec_font
                            break

                    if is_decorative and font_name not in flagged_fonts:
                        flagged_fonts.add(font_name)
                        preview = text[:60] + "..." if len(text) > 60 else text

                        issue = AccessibilityIssue(
                            issue_type="DECORATIVE_FONT",
                            description=f"Font '{font_name}' is decorative/hard to read (use Arial, Calibri, or Times New Roman): \"{preview}\"",
                            location=para_info.location,
                            para_info=para_info
                        )
                        issues.append(issue)
                        break  # Only flag paragraph once

        return issues

    def check_inconsistent_fonts(self) -> List[AccessibilityIssue]:
        """Identify inconsistent font families across the document"""
        issues = []

        # Track font usage across the document
        font_usage = {}  # font_name -> count
        font_locations = {}  # font_name -> list of locations

        # Collect all fonts used
        for para_info in self.all_paragraphs:
            para = para_info.paragraph
            text = para.text.strip()

            # Skip empty paragraphs
            if not text:
                continue

            # Skip headings (they may legitimately use different fonts)
            if para.style.name.startswith('Heading'):
                continue

            # Check each run for font
            for run in para.runs:
                run_text = run.text.strip()
                if not run_text:
                    continue

                font_name = run.font.name
                if font_name:
                    # Count this font
                    font_usage[font_name] = font_usage.get(font_name, 0) + 1

                    # Track location
                    if font_name not in font_locations:
                        font_locations[font_name] = []
                    if para_info.location not in font_locations[font_name]:
                        font_locations[font_name].append(para_info.location)

        # Identify the primary font (most commonly used)
        if not font_usage:
            return issues

        primary_font = max(font_usage, key=font_usage.get)
        primary_count = font_usage[primary_font]
        total_count = sum(font_usage.values())

        # Check for font inconsistency
        # If there are multiple fonts and the primary font isn't dominant (>80%), flag it
        if len(font_usage) > 2:
            # Multiple fonts detected
            primary_ratio = primary_count / total_count if total_count > 0 else 0

            # List secondary fonts (fonts that aren't the primary)
            secondary_fonts = {font: count for font, count in font_usage.items() if font != primary_font}

            # Only flag if there are significant secondary fonts (not just a few instances)
            significant_secondary = {font: count for font, count in secondary_fonts.items()
                                    if count > total_count * 0.1}  # More than 10% usage

            if significant_secondary:
                font_summary = []
                font_summary.append(f"{primary_font} ({primary_count} instances, {primary_ratio*100:.1f}%)")
                for font, count in sorted(significant_secondary.items(), key=lambda x: x[1], reverse=True):
                    ratio = count / total_count * 100
                    font_summary.append(f"{font} ({count} instances, {ratio:.1f}%)")

                issue = AccessibilityIssue(
                    issue_type="INCONSISTENT_FONTS",
                    description=f"Document uses {len(font_usage)} different fonts. Consider using a single font for consistency. Fonts: {'; '.join(font_summary[:3])}{'...' if len(font_summary) > 3 else ''}",
                    location="Document-wide"
                )
                issues.append(issue)

        return issues

    def check_color_contrast(self) -> List[AccessibilityIssue]:
        """Check for insufficient color contrast based on WCAG thresholds"""
        issues = []

        for para_info in self.all_paragraphs:
            para = para_info.paragraph
            text = para.text.strip()

            # Skip empty paragraphs
            if not text:
                continue

            # Check each run for color contrast
            for run in para.runs:
                run_text = run.text.strip()
                if not run_text:
                    continue

                # Get text color (foreground)
                text_color = None
                try:
                    if run.font.color.rgb:
                        text_color = (run.font.color.rgb[0], run.font.color.rgb[1], run.font.color.rgb[2])
                    elif run.font.color.theme_color:
                        # Default to black for theme colors (conservative)
                        text_color = (0, 0, 0)
                    else:
                        # No explicit color set, assume black (default)
                        text_color = (0, 0, 0)
                except:
                    # If we can't get color, assume black
                    text_color = (0, 0, 0)

                # Get background color - check paragraph shading
                bg_color = None
                try:
                    # Check paragraph shading
                    pPr = para._element.get_or_add_pPr()
                    shd = pPr.find(qn('w:shd'))
                    if shd is not None:
                        fill = shd.get(qn('w:fill'))
                        if fill and fill != 'auto' and len(fill) == 6:
                            # Convert hex to RGB
                            bg_color = tuple(int(fill[i:i+2], 16) for i in (0, 2, 4))
                except:
                    pass

                # If no paragraph background, assume white (default)
                if bg_color is None:
                    bg_color = (255, 255, 255)

                # Calculate contrast ratio
                if text_color and bg_color:
                    contrast_ratio = calculate_contrast_ratio(text_color, bg_color)

                    # Determine if text is "large" (18pt+ or 14pt+ bold)
                    is_large_text = False
                    font_size = run.font.size
                    if font_size:
                        size_pt = font_size.pt
                        if run.bold and size_pt >= 14:
                            is_large_text = True
                        elif size_pt >= 18:
                            is_large_text = True

                    # WCAG AA thresholds
                    required_ratio = 3.0 if is_large_text else 4.5

                    if contrast_ratio < required_ratio:
                        preview = text[:60] + "..." if len(text) > 60 else text
                        text_color_hex = '#{:02x}{:02x}{:02x}'.format(*text_color)
                        bg_color_hex = '#{:02x}{:02x}{:02x}'.format(*bg_color)

                        issue = AccessibilityIssue(
                            issue_type="LOW_CONTRAST",
                            description=f"Insufficient color contrast {contrast_ratio:.2f}:1 (minimum {required_ratio}:1). Text: {text_color_hex}, Background: {bg_color_hex}. \"{preview}\"",
                            location=para_info.location,
                            para_info=para_info
                        )
                        issues.append(issue)
                        break  # Only flag paragraph once

        return issues

    def check_color_only_meaning(self) -> List[AccessibilityIssue]:
        """Flag color used as the sole means of conveying meaning"""
        issues = []

        # Track colored text runs to identify patterns
        colored_runs = []

        for para_info in self.all_paragraphs:
            para = para_info.paragraph
            text = para.text.strip()

            if not text:
                continue

            # Check for runs with different colors but no other distinguishing features
            run_colors = []
            for run in para.runs:
                run_text = run.text.strip()
                if not run_text:
                    continue

                try:
                    # Check if run has explicit color
                    has_color = run.font.color.rgb is not None or run.font.color.theme_color is not None

                    # Check for other formatting
                    has_bold = run.bold
                    has_italic = run.italic
                    has_underline = run.font.underline

                    if has_color and not (has_bold or has_italic or has_underline):
                        # Color is used without other formatting
                        color_rgb = None
                        try:
                            if run.font.color.rgb:
                                color_rgb = (run.font.color.rgb[0], run.font.color.rgb[1], run.font.color.rgb[2])
                        except:
                            pass

                        run_colors.append({
                            'text': run_text,
                            'color': color_rgb,
                            'para_info': para_info
                        })
                except:
                    pass

            # If we found multiple colored runs in the same paragraph, flag it
            if len(run_colors) >= 2:
                # Check if they have different colors
                unique_colors = set()
                for rc in run_colors:
                    if rc['color']:
                        unique_colors.add(rc['color'])

                if len(unique_colors) > 1:
                    preview = text[:60] + "..." if len(text) > 60 else text
                    issue = AccessibilityIssue(
                        issue_type="COLOR_ONLY_MEANING",
                        description=f"Color appears to be the only way to distinguish text (add bold, italic, or other formatting): \"{preview}\"",
                        location=para_info.location,
                        para_info=para_info
                    )
                    issues.append(issue)

        # Check tables for color-coded cells
        for table_idx, table in enumerate(self.target_doc.tables):
            # Track cell background colors
            cell_colors = {}

            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    try:
                        tc = cell._element
                        tcPr = tc.get_or_add_tcPr()
                        shd = tcPr.find(qn('w:shd'))
                        if shd is not None:
                            fill = shd.get(qn('w:fill'))
                            if fill and fill != 'auto' and fill != 'FFFFFF':
                                # This cell has a background color
                                cell_text = cell.text.strip()
                                if fill not in cell_colors:
                                    cell_colors[fill] = []
                                cell_colors[fill].append({
                                    'row': row_idx,
                                    'col': cell_idx,
                                    'text': cell_text[:30]
                                })
                    except:
                        pass

            # If table has multiple background colors, it might be using color for meaning
            if len(cell_colors) > 1:
                issue = AccessibilityIssue(
                    issue_type="COLOR_CODED_TABLE",
                    description=f"Table {table_idx + 1} uses {len(cell_colors)} different background colors. Ensure color is not the only way to convey meaning (use icons, text labels, or patterns).",
                    location=f"Table {table_idx + 1}",
                    table_index=table_idx
                )
                issues.append(issue)

        return issues

    def check_text_over_backgrounds(self) -> List[AccessibilityIssue]:
        """Identify text placed over backgrounds/images without readable contrast"""
        issues = []

        for para_info in self.all_paragraphs:
            para = para_info.paragraph
            text = para.text.strip()

            # Skip empty paragraphs
            if not text:
                continue

            # Check for paragraph background/shading
            has_background = False
            bg_color = None

            try:
                pPr = para._element.get_or_add_pPr()
                shd = pPr.find(qn('w:shd'))
                if shd is not None:
                    fill = shd.get(qn('w:fill'))
                    if fill and fill != 'auto':
                        has_background = True
                        # Convert hex to RGB
                        if len(fill) == 6:
                            bg_color = tuple(int(fill[i:i+2], 16) for i in (0, 2, 4))
            except:
                pass

            if has_background:
                # Check contrast of all text runs against this background
                for run in para.runs:
                    run_text = run.text.strip()
                    if not run_text:
                        continue

                    # Get text color
                    text_color = None
                    try:
                        if run.font.color.rgb:
                            text_color = (run.font.color.rgb[0], run.font.color.rgb[1], run.font.color.rgb[2])
                        else:
                            # Default to black
                            text_color = (0, 0, 0)
                    except:
                        text_color = (0, 0, 0)

                    if text_color and bg_color:
                        contrast_ratio = calculate_contrast_ratio(text_color, bg_color)

                        # Check if text is large
                        is_large_text = False
                        if run.font.size:
                            size_pt = run.font.size.pt
                            if run.bold and size_pt >= 14:
                                is_large_text = True
                            elif size_pt >= 18:
                                is_large_text = True

                        required_ratio = 3.0 if is_large_text else 4.5

                        if contrast_ratio < required_ratio:
                            preview = text[:60] + "..." if len(text) > 60 else text
                            bg_color_hex = '#{:02x}{:02x}{:02x}'.format(*bg_color)

                            issue = AccessibilityIssue(
                                issue_type="TEXT_OVER_BACKGROUND",
                                description=f"Text over colored background (#{bg_color_hex}) has low contrast {contrast_ratio:.2f}:1 (minimum {required_ratio}:1): \"{preview}\"",
                                location=para_info.location,
                                para_info=para_info
                            )
                            issues.append(issue)
                            break  # Only flag paragraph once

        # Check table cells with backgrounds
        for table_idx, table in enumerate(self.target_doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    # Get cell background color
                    bg_color = None
                    try:
                        tc = cell._element
                        tcPr = tc.get_or_add_tcPr()
                        shd = tcPr.find(qn('w:shd'))
                        if shd is not None:
                            fill = shd.get(qn('w:fill'))
                            if fill and fill != 'auto' and len(fill) == 6:
                                bg_color = tuple(int(fill[i:i+2], 16) for i in (0, 2, 4))
                    except:
                        pass

                    if bg_color:
                        # Check text in this cell
                        for para in cell.paragraphs:
                            para_text = para.text.strip()
                            if not para_text:
                                continue

                            for run in para.runs:
                                run_text = run.text.strip()
                                if not run_text:
                                    continue

                                # Get text color
                                text_color = None
                                try:
                                    if run.font.color.rgb:
                                        text_color = (run.font.color.rgb[0], run.font.color.rgb[1], run.font.color.rgb[2])
                                    else:
                                        text_color = (0, 0, 0)
                                except:
                                    text_color = (0, 0, 0)

                                if text_color:
                                    contrast_ratio = calculate_contrast_ratio(text_color, bg_color)

                                    # Check if large text
                                    is_large_text = False
                                    if run.font.size:
                                        size_pt = run.font.size.pt
                                        if run.bold and size_pt >= 14:
                                            is_large_text = True
                                        elif size_pt >= 18:
                                            is_large_text = True

                                    required_ratio = 3.0 if is_large_text else 4.5

                                    if contrast_ratio < required_ratio:
                                        preview = para_text[:60] + "..." if len(para_text) > 60 else para_text
                                        bg_color_hex = '#{:02x}{:02x}{:02x}'.format(*bg_color)

                                        issue = AccessibilityIssue(
                                            issue_type="TEXT_OVER_BACKGROUND",
                                            description=f"Text in table cell over colored background (#{bg_color_hex}) has low contrast {contrast_ratio:.2f}:1 (minimum {required_ratio}:1): \"{preview}\"",
                                            location=f"Table {table_idx + 1}, Row {row_idx + 1}, Cell {cell_idx + 1}",
                                            table_index=table_idx
                                        )
                                        issues.append(issue)
                                        break  # Only flag cell once

        return issues

    def check_empty_table_rows_columns(self) -> List[AccessibilityIssue]:
        """Detect empty rows/columns in tables used for spacing"""
        issues = []

        for table_idx, table in enumerate(self.target_doc.tables):
            # Check for completely empty rows
            for row_idx, row in enumerate(table.rows):
                all_empty = True
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        all_empty = False
                        break

                if all_empty:
                    issue = AccessibilityIssue(
                        issue_type="EMPTY_TABLE_ROW",
                        description=f"Table {table_idx + 1}, Row {row_idx + 1} is completely empty (likely used for spacing)",
                        location=f"Table {table_idx + 1}, Row {row_idx + 1}",
                        table_index=table_idx
                    )
                    issues.append(issue)

            # Check for completely empty columns
            num_rows = len(table.rows)
            num_cols = len(table.columns) if table.rows else 0

            if num_rows > 0 and num_cols > 0:
                for col_idx in range(num_cols):
                    all_empty = True
                    for row in table.rows:
                        if col_idx < len(row.cells):
                            cell_text = row.cells[col_idx].text.strip()
                            if cell_text:
                                all_empty = False
                                break

                    if all_empty:
                        issue = AccessibilityIssue(
                            issue_type="EMPTY_TABLE_COLUMN",
                            description=f"Table {table_idx + 1}, Column {col_idx + 1} is completely empty (likely used for spacing)",
                            location=f"Table {table_idx + 1}, Column {col_idx + 1}",
                            table_index=table_idx
                        )
                        issues.append(issue)

        return issues

    def check_table_scope_declarations(self) -> List[AccessibilityIssue]:
        """Flag tables missing proper scope declarations on header cells"""
        issues = []

        for table_idx, table in enumerate(self.target_doc.tables):
            # Check if table has a header row first
            if not self._table_has_header_row(table):
                continue  # Skip tables without headers

            # Check first row for scope attributes
            first_row = table.rows[0]
            missing_scope_count = 0

            for cell_idx, cell in enumerate(first_row.cells):
                # Check if cell has proper scope declaration in XML
                try:
                    tc = cell._element
                    # Look for scope in table cell properties
                    # In Word XML, header cells should be marked with specific properties
                    # Check if this is a proper table header cell (th) or just regular cell (td)

                    # Word doesn't use HTML-style scope attributes, but we can check
                    # if the cell is properly marked as a header cell semantically
                    # For now, we'll flag tables that appear to have headers but may lack
                    # proper semantic markup for screen readers

                    # Since Word's accessibility is different from HTML, we'll check
                    # if the table has its first row marked as a "repeat as header row"
                    # which is Word's way of indicating header rows
                    missing_scope_count += 1
                except:
                    missing_scope_count += 1

            # If table has a header row but cells don't have proper scope info
            if missing_scope_count > 0:
                issue = AccessibilityIssue(
                    issue_type="TABLE_MISSING_SCOPE",
                    description=f"Table {table_idx + 1} has header row but may lack proper scope declarations for screen readers (consider using Table Tools > Layout > Repeat Header Rows)",
                    location=f"Table {table_idx + 1}",
                    table_index=table_idx
                )
                issues.append(issue)

        return issues

    def check_table_captions(self) -> List[AccessibilityIssue]:
        """Detect missing captions or table descriptions"""
        issues = []

        for table_idx, table in enumerate(self.target_doc.tables):
            # Check if table has a caption or title
            # In Word, captions are typically added as paragraphs before/after the table
            # or in the table properties (Alt Text)

            has_caption = False

            # Method 1: Check for Alt Text (Title/Description)
            try:
                # Access table properties to check for title/description
                tbl = table._element
                tblPr = tbl.find(qn('w:tblPr'))
                if tblPr is not None:
                    # Look for table description in Alt Text
                    tblDescription = tblPr.find(qn('w:tblDescription'))
                    tblTitle = tblPr.find(qn('w:tblCaption'))

                    if tblDescription is not None:
                        desc_val = tblDescription.get(qn('w:val'))
                        if desc_val and desc_val.strip():
                            has_caption = True

                    if tblTitle is not None:
                        title_val = tblTitle.get(qn('w:val'))
                        if title_val and title_val.strip():
                            has_caption = True
            except:
                pass

            if not has_caption:
                # Get table preview for context
                preview = ""
                if table.rows:
                    preview = " | ".join([cell.text.strip()[:15] for cell in table.rows[0].cells[:3]])

                issue = AccessibilityIssue(
                    issue_type="TABLE_MISSING_CAPTION",
                    description=f"Table {table_idx + 1} is missing a caption or description (Right-click table > Table Properties > Alt Text). Preview: \"{preview}...\"",
                    location=f"Table {table_idx + 1}",
                    table_index=table_idx
                )
                issues.append(issue)

        return issues

    def check_table_merged_cells(self) -> List[AccessibilityIssue]:
        """Detect merged or split cells that may impact reading order"""
        issues = []

        for table_idx, table in enumerate(self.target_doc.tables):
            merged_cells_found = False
            merged_locations = []

            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    try:
                        tc = cell._element
                        tcPr = tc.find(qn('w:tcPr'))

                        if tcPr is not None:
                            # Check for horizontal merge (gridSpan)
                            gridSpan = tcPr.find(qn('w:gridSpan'))
                            if gridSpan is not None:
                                span_val = gridSpan.get(qn('w:val'))
                                if span_val and int(span_val) > 1:
                                    merged_cells_found = True
                                    merged_locations.append(f"Row {row_idx + 1}, Col {cell_idx + 1}")

                            # Check for vertical merge (vMerge)
                            vMerge = tcPr.find(qn('w:vMerge'))
                            if vMerge is not None:
                                merged_cells_found = True
                                if f"Row {row_idx + 1}, Col {cell_idx + 1}" not in merged_locations:
                                    merged_locations.append(f"Row {row_idx + 1}, Col {cell_idx + 1}")
                    except:
                        pass

            if merged_cells_found:
                locations_str = "; ".join(merged_locations[:3])  # Show first 3 locations
                if len(merged_locations) > 3:
                    locations_str += f" (and {len(merged_locations) - 3} more)"

                issue = AccessibilityIssue(
                    issue_type="TABLE_MERGED_CELLS",
                    description=f"Table {table_idx + 1} contains merged cells which may impact screen reader navigation. Locations: {locations_str}",
                    location=f"Table {table_idx + 1}",
                    table_index=table_idx
                )
                issues.append(issue)

        return issues

    def check_table_numeric_alignment(self) -> List[AccessibilityIssue]:
        """Identify non-left-aligned numeric data inconsistently applied"""
        issues = []

        for table_idx, table in enumerate(self.target_doc.tables):
            # Track numeric columns and their alignment
            num_rows = len(table.rows)
            num_cols = len(table.columns) if table.rows else 0

            if num_rows == 0 or num_cols == 0:
                continue

            # Skip header row if present
            start_row = 1 if self._table_has_header_row(table) else 0

            # Check each column for numeric data
            for col_idx in range(num_cols):
                numeric_count = 0
                total_count = 0
                alignments = []

                for row_idx in range(start_row, num_rows):
                    if col_idx < len(table.rows[row_idx].cells):
                        cell = table.rows[row_idx].cells[col_idx]
                        cell_text = cell.text.strip()

                        if cell_text:
                            total_count += 1

                            # Check if cell contains numeric data
                            # Remove common formatting (currency symbols, commas, %)
                            cleaned = cell_text.replace('$', '').replace(',', '').replace('%', '').strip()
                            try:
                                float(cleaned)
                                numeric_count += 1

                                # Check alignment of this cell
                                try:
                                    for para in cell.paragraphs:
                                        if para.text.strip():
                                            alignment = para.alignment
                                            if alignment is not None:
                                                alignments.append(str(alignment))
                                            break  # Only check first paragraph
                                except:
                                    pass
                            except ValueError:
                                pass

                # If column is mostly numeric (>70%) and has inconsistent alignment
                if total_count > 0 and numeric_count / total_count > 0.7:
                    # Check if alignments are inconsistent
                    if len(set(alignments)) > 1:
                        issue = AccessibilityIssue(
                            issue_type="TABLE_INCONSISTENT_NUMERIC_ALIGNMENT",
                            description=f"Table {table_idx + 1}, Column {col_idx + 1} contains numeric data with inconsistent alignment (found {len(set(alignments))} different alignments)",
                            location=f"Table {table_idx + 1}, Column {col_idx + 1}",
                            table_index=table_idx
                        )
                        issues.append(issue)

        return issues

    def check_table_reading_order(self) -> List[AccessibilityIssue]:
        """Flag inconsistent or illogical table reading order"""
        issues = []

        for table_idx, table in enumerate(self.target_doc.tables):
            # Check for tables with complex structures that might confuse reading order

            # Factor 1: Check if table has both merged cells AND multiple header patterns
            has_merged = False
            try:
                for row in table.rows:
                    for cell in row.cells:
                        tc = cell._element
                        tcPr = tc.find(qn('w:tcPr'))
                        if tcPr is not None:
                            if tcPr.find(qn('w:gridSpan')) is not None or tcPr.find(qn('w:vMerge')) is not None:
                                has_merged = True
                                break
                    if has_merged:
                        break
            except:
                pass

            # Factor 2: Check if table has multiple "header-like" rows (bold rows beyond first)
            header_like_rows = []
            for row_idx, row in enumerate(table.rows):
                bold_count = 0
                total_runs = 0

                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            total_runs += 1
                            if run.bold:
                                bold_count += 1

                if total_runs > 0 and bold_count / total_runs > 0.7:
                    header_like_rows.append(row_idx)

            # If table has merged cells AND multiple header rows, reading order may be confusing
            if has_merged and len(header_like_rows) > 1:
                issue = AccessibilityIssue(
                    issue_type="TABLE_COMPLEX_READING_ORDER",
                    description=f"Table {table_idx + 1} has complex structure (merged cells + multiple header-like rows at positions {header_like_rows}) which may create confusing reading order",
                    location=f"Table {table_idx + 1}",
                    table_index=table_idx
                )
                issues.append(issue)

        return issues

    def check_table_color_coded_meaning(self) -> List[AccessibilityIssue]:
        """Detect color-coded meaning in cells without textual explanation"""
        issues = []

        for table_idx, table in enumerate(self.target_doc.tables):
            # Track cells with background colors
            colored_cells = {}  # color -> list of (row, col, text)

            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()

                    # Check for cell background color
                    try:
                        tc = cell._element
                        tcPr = tc.find(qn('w:tcPr'))
                        if tcPr is not None:
                            shd = tcPr.find(qn('w:shd'))
                            if shd is not None:
                                fill = shd.get(qn('w:fill'))
                                # Ignore white/auto backgrounds
                                if fill and fill != 'auto' and fill.upper() not in ['FFFFFF', 'FFFF', 'FFF']:
                                    if fill not in colored_cells:
                                        colored_cells[fill] = []
                                    colored_cells[fill].append((row_idx, cell_idx, cell_text[:30]))
                    except:
                        pass

            # If table has multiple distinct colors used across cells
            if len(colored_cells) >= 2:
                # Check if there's explanatory text (legend) in the table or nearby
                # Simple heuristic: if cells with same color have very similar content,
                # color might be conveying meaning

                potentially_meaningful_colors = []
                for color, cells in colored_cells.items():
                    if len(cells) >= 2:  # Color used multiple times
                        # Check if cells with this color have varying content
                        texts = [c[2] for c in cells]
                        unique_texts = len(set(texts))
                        if unique_texts > 1:  # Different content, same color
                            potentially_meaningful_colors.append(color)

                if len(potentially_meaningful_colors) >= 2:
                    issue = AccessibilityIssue(
                        issue_type="TABLE_COLOR_CODED_MEANING",
                        description=f"Table {table_idx + 1} uses {len(potentially_meaningful_colors)} different background colors which may convey meaning. Ensure there's a legend or textual explanation.",
                        location=f"Table {table_idx + 1}",
                        table_index=table_idx
                    )
                    issues.append(issue)

        return issues

    def check_table_embedded_images(self) -> List[AccessibilityIssue]:
        """Identify images embedded inside table cells when text should be used"""
        issues = []

        for table_idx, table in enumerate(self.target_doc.tables):
            cells_with_images = []

            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    # Check if cell contains images
                    has_image = False
                    cell_text = cell.text.strip()

                    for para in cell.paragraphs:
                        # Check for inline shapes (images)
                        try:
                            # Look for drawing elements in the paragraph
                            para_element = para._element
                            drawings = para_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                            if drawings:
                                has_image = True
                                break

                            # Also check for pict elements (older image format)
                            picts = para_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict')
                            if picts:
                                has_image = True
                                break
                        except:
                            pass

                    if has_image:
                        # Flag if image is in cell with little/no text
                        if len(cell_text) < 10:  # Arbitrary threshold
                            cells_with_images.append(f"Row {row_idx + 1}, Col {cell_idx + 1}")

            if cells_with_images:
                locations_str = "; ".join(cells_with_images[:3])
                if len(cells_with_images) > 3:
                    locations_str += f" (and {len(cells_with_images) - 3} more)"

                issue = AccessibilityIssue(
                    issue_type="TABLE_EMBEDDED_IMAGES",
                    description=f"Table {table_idx + 1} contains images in cells with minimal text. Consider using text instead or ensure images have alt text. Locations: {locations_str}",
                    location=f"Table {table_idx + 1}",
                    table_index=table_idx
                )
                issues.append(issue)

        return issues

    def check_document_metadata(self) -> List[AccessibilityIssue]:
        """Check for missing document title"""
        issues = []

        # Check document title
        core_properties = self.target_doc.core_properties
        title = core_properties.title

        if not title or title.strip() == "":
            issue = AccessibilityIssue(
                issue_type="MISSING_TITLE",
                description="Document is missing a title in metadata (File > Info > Properties)",
                location="Document metadata"
            )
            issues.append(issue)

        return issues

    def check_document_language(self) -> List[AccessibilityIssue]:
        """Check for missing or incorrect document language setting"""
        issues = []

        # Check if document has a default language set
        has_language = False
        document_language = None

        try:
            # Check the document's default paragraph style for language
            # Language is typically set in the default style or document settings
            for style in self.target_doc.styles:
                if style.name == 'Normal' or style.name == 'Default Paragraph Font':
                    try:
                        # Try to access the style's element to check for language
                        style_element = style.element
                        # Look for language setting in the style
                        rPr = style_element.find(qn('w:rPr'))
                        if rPr is not None:
                            lang = rPr.find(qn('w:lang'))
                            if lang is not None:
                                val = lang.get(qn('w:val'))
                                if val:
                                    has_language = True
                                    document_language = val
                                    break
                    except:
                        pass
        except:
            pass

        # Also check if any paragraphs have explicit language settings
        if not has_language:
            # Sample first few paragraphs to check for language
            for para_info in self.all_paragraphs[:10]:
                para = para_info.paragraph
                if para.runs:
                    for run in para.runs:
                        try:
                            # Check run's language
                            rPr = run._element.get_or_add_rPr()
                            lang = rPr.find(qn('w:lang'))
                            if lang is not None:
                                val = lang.get(qn('w:val'))
                                if val:
                                    has_language = True
                                    document_language = val
                                    break
                        except:
                            pass
                if has_language:
                    break

        if not has_language:
            issue = AccessibilityIssue(
                issue_type="MISSING_LANGUAGE",
                description="Document is missing language setting. Set document language in Word (Review > Language > Set Proofing Language) for screen reader accessibility.",
                location="Document metadata"
            )
            issues.append(issue)
        else:
            # Validate that the language code is reasonable
            # Common language codes: en-US, en-GB, es-ES, fr-FR, etc.
            if document_language and len(document_language) < 2:
                issue = AccessibilityIssue(
                    issue_type="INVALID_LANGUAGE",
                    description=f"Document language setting '{document_language}' appears invalid. Use standard language codes (e.g., en-US, es-ES).",
                    location="Document metadata"
                )
                issues.append(issue)

        return issues

    def check_multilingual_content(self) -> List[AccessibilityIssue]:
        """Identify multiple languages without proper tagging"""
        issues = []

        # Track different languages used in the document
        languages_found = set()
        language_locations = {}  # language -> list of locations

        for para_info in self.all_paragraphs:
            para = para_info.paragraph
            text = para.text.strip()

            if not text:
                continue

            # Check each run for language tags
            for run in para.runs:
                run_text = run.text.strip()
                if not run_text:
                    continue

                try:
                    # Check if run has explicit language setting
                    rPr = run._element.get_or_add_rPr()
                    lang = rPr.find(qn('w:lang'))

                    lang_code = None
                    if lang is not None:
                        val = lang.get(qn('w:val'))
                        if val:
                            lang_code = val
                            languages_found.add(lang_code)

                            if lang_code not in language_locations:
                                language_locations[lang_code] = []
                            if para_info.location not in language_locations[lang_code]:
                                language_locations[lang_code].append(para_info.location)
                except:
                    pass

        # If we found multiple languages, check if they're properly tagged
        if len(languages_found) > 1:
            # Multiple languages detected - this is fine if properly tagged
            # Just inform the user
            lang_summary = []
            for lang in sorted(languages_found):
                count = len(language_locations.get(lang, []))
                lang_summary.append(f"{lang} ({count} locations)")

            issue = AccessibilityIssue(
                issue_type="MULTIPLE_LANGUAGES",
                description=f"Document contains {len(languages_found)} different language tags: {', '.join(lang_summary)}. Verify all content is properly tagged with correct language for screen reader pronunciation.",
                location="Document-wide"
            )
            issues.append(issue)

        # Detect common non-English phrases that might need language tagging
        # This is a heuristic check for potential multilingual content
        non_english_patterns = {
            'spanish': [r'\bseñor\b', r'\bseñora\b', r'\bhola\b', r'\bgracias\b', r'\bpor favor\b',
                       r'\b¿\w+\?', r'\b¡\w+!', r'\baño\b', r'\bniño\b', r'\bniña\b'],
            'french': [r'\bmonsieur\b', r'\bmadame\b', r'\bbonjour\b', r'\bmerci\b',
                      r'\bs\'il vous plaît\b', r'\bça\b', r'\bà\b', r'\bélève\b'],
            'german': [r'\bherr\b', r'\bfrau\b', r'\bguten tag\b', r'\bdanke\b',
                      r'\bbitte\b', r'\bstraße\b', r'\bschön\b', r'\büber\b'],
        }

        potential_multilingual = []

        for para_info in self.all_paragraphs:
            para = para_info.paragraph
            text = para.text.lower()

            if not text or len(text) < 10:
                continue

            # Check for non-English patterns
            for language, patterns in non_english_patterns.items():
                for pattern in patterns:
                    if re.search(pattern, text, re.IGNORECASE):
                        # Check if this paragraph has proper language tagging
                        has_proper_tag = False
                        for run in para.runs:
                            try:
                                rPr = run._element.get_or_add_rPr()
                                lang = rPr.find(qn('w:lang'))
                                if lang is not None:
                                    val = lang.get(qn('w:val'))
                                    if val and not val.startswith('en-'):
                                        has_proper_tag = True
                                        break
                            except:
                                pass

                        if not has_proper_tag:
                            preview = text[:60] + "..." if len(text) > 60 else text
                            potential_multilingual.append({
                                'location': para_info.location,
                                'language': language,
                                'text': preview,
                                'para_info': para_info
                            })
                            break  # Only flag once per paragraph
                    if potential_multilingual and potential_multilingual[-1].get('location') == para_info.location:
                        break

        if potential_multilingual:
            # Group by language
            by_language = {}
            for item in potential_multilingual:
                lang = item['language']
                if lang not in by_language:
                    by_language[lang] = []
                by_language[lang].append(item)

            for language, items in by_language.items():
                issue = AccessibilityIssue(
                    issue_type="UNTAGGED_LANGUAGE",
                    description=f"Found {len(items)} paragraph(s) with potential {language.title()} content that may need language tagging. Example: \"{items[0]['text']}\" at {items[0]['location']}",
                    location="Document-wide",
                    para_info=items[0]['para_info']
                )
                issues.append(issue)

        return issues

    def check_underline_non_links(self) -> List[AccessibilityIssue]:
        """Detect underlined text that is not a hyperlink (confusing for navigation)"""
        issues = []

        for para_info in self.all_paragraphs:
            para = para_info.paragraph
            text = para.text.strip()

            if not text:
                continue

            # Check each run for underline
            for run in para.runs:
                run_text = run.text.strip()

                # Skip if no text
                if not run_text:
                    continue

                # Check if underlined
                if run.font.underline:
                    # Check if it's a hyperlink by checking for hyperlink relationship
                    # In python-docx, hyperlinks are tricky to detect
                    # We'll use a heuristic: if it contains http:// or www., it's likely a link
                    is_likely_link = 'http://' in run_text or 'https://' in run_text or 'www.' in run_text

                    # Also check if the run has a hyperlink style
                    # Hyperlinks often have specific formatting
                    has_link_color = False
                    if run.font.color.rgb:
                        # Common link colors are blue (0, 0, 255) or similar
                        rgb = run.font.color.rgb
                        # Check if it's bluish
                        if rgb.blue > 200 and rgb.red < 100 and rgb.green < 100:
                            has_link_color = True

                    if not is_likely_link and not has_link_color:
                        preview = text[:60] + "..." if len(text) > 60 else text

                        issue = AccessibilityIssue(
                            issue_type="UNDERLINE_NON_LINK",
                            description=f"Underlined text that is not a hyperlink: \"{preview}\"",
                            location=para_info.location,
                            para_info=para_info
                        )
                        issues.append(issue)
                        break  # Only flag paragraph once

        return issues

    def check_line_spacing(self) -> List[AccessibilityIssue]:
        """Check for line spacing below 1.15 (WCAG recommendation)"""
        issues = []

        for para_info in self.all_paragraphs:
            para = para_info.paragraph
            text = para.text.strip()

            # Skip empty paragraphs
            if not text:
                continue

            # Get line spacing
            # In python-docx, line spacing is in paragraph_format.line_spacing
            # Can be a float (multiple of single spacing) or None (uses default)
            if para.paragraph_format.line_spacing is not None:
                line_spacing = para.paragraph_format.line_spacing

                # Line spacing can be specified as a float (1.0, 1.15, etc.)
                # or as a Pt value - we need to handle both
                if hasattr(line_spacing, '__float__'):
                    spacing_value = float(line_spacing)
                    if spacing_value < 1.15:
                        preview = text[:60] + "..." if len(text) > 60 else text

                        issue = AccessibilityIssue(
                            issue_type="LOW_LINE_SPACING",
                            description=f"Line spacing {spacing_value:.2f} (minimum 1.15): \"{preview}\"",
                            location=para_info.location,
                            para_info=para_info
                        )
                        issues.append(issue)

        return issues

    def check_full_justification(self) -> List[AccessibilityIssue]:
        """Check for full text justification (hard to read, especially for dyslexic readers)"""
        issues = []

        for para_info in self.all_paragraphs:
            para = para_info.paragraph
            text = para.text.strip()

            # Skip empty paragraphs
            if not text:
                continue

            # Check alignment - WD_ALIGN_PARAGRAPH.JUSTIFY = 3
            if para.paragraph_format.alignment == 3:  # JUSTIFY
                preview = text[:60] + "..." if len(text) > 60 else text

                issue = AccessibilityIssue(
                    issue_type="FULL_JUSTIFICATION",
                    description=f"Text is fully justified (harder to read): \"{preview}\"",
                    location=para_info.location,
                    para_info=para_info
                )
                issues.append(issue)

        return issues

    def check_all_caps_blocks(self) -> List[AccessibilityIssue]:
        """Check for large blocks of ALL CAPS text (hard to read)"""
        issues = []

        for para_info in self.all_paragraphs:
            para = para_info.paragraph
            text = para.text.strip()

            # Skip empty paragraphs
            if not text:
                continue

            # Skip very short text (headings already handled elsewhere)
            # We're looking for BLOCKS of caps text, not short headings
            if len(text) < 50:  # Less than 50 chars is likely a heading/label
                continue

            # Check if most of the text is uppercase
            # Count letters only (ignore numbers, punctuation, spaces)
            letters = [c for c in text if c.isalpha()]
            if len(letters) > 0:
                uppercase_letters = [c for c in letters if c.isupper()]
                caps_ratio = len(uppercase_letters) / len(letters)

                # If more than 80% uppercase and it's a substantial block, flag it
                if caps_ratio > 0.8:
                    preview = text[:60] + "..." if len(text) > 60 else text

                    issue = AccessibilityIssue(
                        issue_type="ALL_CAPS_BLOCK",
                        description=f"Large block of ALL CAPS text (hard to read): \"{preview}\"",
                        location=para_info.location,
                        para_info=para_info
                    )
                    issues.append(issue)

        return issues

    def check_non_descriptive_links(self) -> List[AccessibilityIssue]:
        """Check for non-descriptive link text like 'click here', 'here', 'read more'"""
        issues = []

        # Common non-descriptive link text patterns
        non_descriptive_patterns = [
            'click here', 'here', 'read more', 'more', 'link', 'this link',
            'click', 'download', 'more info', 'more information', 'learn more',
            'see more', 'view more', 'details', 'continue', 'next', 'previous'
        ]

        for para_info in self.all_paragraphs:
            para = para_info.paragraph

            # Check hyperlinks in this paragraph
            # In python-docx, hyperlinks are tricky - we'll check run text
            for run in para.runs:
                run_text = run.text.strip().lower()

                # Skip if no text
                if not run_text:
                    continue

                # Check if underlined or has link color (heuristic for links)
                is_likely_link = False
                if run.font.underline:
                    is_likely_link = True
                elif run.font.color.rgb:
                    try:
                        rgb = run.font.color.rgb
                        # RGBColor is a tuple-like (r, g, b)
                        # Check if it's bluish (common link color)
                        if len(rgb) >= 3 and rgb[2] > 200 and rgb[0] < 100 and rgb[1] < 100:
                            is_likely_link = True
                    except:
                        pass  # If we can't check color, skip it

                # If it looks like a link, check if text is descriptive
                if is_likely_link:
                    # Check against non-descriptive patterns
                    if run_text in non_descriptive_patterns:
                        preview = para.text[:60] + "..." if len(para.text) > 60 else para.text

                        issue = AccessibilityIssue(
                            issue_type="NON_DESCRIPTIVE_LINK",
                            description=f"Link text '{run_text}' is not descriptive: \"{preview}\"",
                            location=para_info.location,
                            para_info=para_info
                        )
                        issues.append(issue)
                        break  # Only flag paragraph once

        return issues

    def check_unstyled_links(self) -> List[AccessibilityIssue]:
        """Identify links styled as normal text (missing underline/color)"""
        issues = []

        for para_info in self.all_paragraphs:
            para = para_info.paragraph

            # Check for hyperlinks in the paragraph XML
            try:
                # Access paragraph element to find hyperlinks
                para_element = para._element
                hyperlinks = para_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink')

                for hyperlink in hyperlinks:
                    # Get the text content of the hyperlink
                    link_text = ''.join(hyperlink.itertext())

                    # Check the formatting of runs within the hyperlink
                    runs_in_link = hyperlink.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')

                    has_underline = False
                    has_distinctive_color = False

                    for run_elem in runs_in_link:
                        # Check for underline
                        rPr = run_elem.find(qn('w:rPr'))
                        if rPr is not None:
                            u = rPr.find(qn('w:u'))
                            if u is not None:
                                has_underline = True

                            # Check for color (non-black)
                            color_elem = rPr.find(qn('w:color'))
                            if color_elem is not None:
                                color_val = color_elem.get(qn('w:val'))
                                # Check if color is not black or auto
                                if color_val and color_val.upper() not in ['000000', '000', 'AUTO', 'FFFFFF']:
                                    has_distinctive_color = True

                    # If link has neither underline nor distinctive color, flag it
                    if not has_underline and not has_distinctive_color:
                        preview = link_text[:40] + "..." if len(link_text) > 40 else link_text

                        issue = AccessibilityIssue(
                            issue_type="UNSTYLED_LINK",
                            description=f"Link styled as normal text (no underline or color): \"{preview}\"",
                            location=para_info.location,
                            para_info=para_info
                        )
                        issues.append(issue)
            except:
                pass  # If we can't parse hyperlinks, skip

        return issues

    def check_long_urls(self) -> List[AccessibilityIssue]:
        """Detect excessively long URLs that should be shortened or use descriptive text"""
        issues = []

        for para_info in self.all_paragraphs:
            para = para_info.paragraph

            # Check for hyperlinks in the paragraph XML
            try:
                para_element = para._element
                hyperlinks = para_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink')

                for hyperlink in hyperlinks:
                    # Get the URL from the relationship ID
                    r_id = hyperlink.get(qn('r:id'))
                    if r_id:
                        try:
                            # Get the relationship to find the actual URL
                            url = para.part.rels[r_id].target_ref

                            # Check if URL is excessively long (>100 characters)
                            if len(url) > 100:
                                link_text = ''.join(hyperlink.itertext())
                                preview = link_text[:40] + "..." if len(link_text) > 40 else link_text

                                issue = AccessibilityIssue(
                                    issue_type="LONG_URL",
                                    description=f"Excessively long URL ({len(url)} chars): \"{preview}\" -> {url[:50]}...",
                                    location=para_info.location,
                                    para_info=para_info
                                )
                                issues.append(issue)
                        except:
                            pass
            except:
                pass

        return issues

    def check_pdf_links(self) -> List[AccessibilityIssue]:
        """Identify links to external PDF files which may not be accessible"""
        issues = []

        for para_info in self.all_paragraphs:
            para = para_info.paragraph

            # Check for hyperlinks in the paragraph XML
            try:
                para_element = para._element
                hyperlinks = para_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink')

                for hyperlink in hyperlinks:
                    # Get the URL from the relationship ID
                    r_id = hyperlink.get(qn('r:id'))
                    if r_id:
                        try:
                            # Get the relationship to find the actual URL
                            url = para.part.rels[r_id].target_ref

                            # Check if URL points to a PDF file
                            if url.lower().endswith('.pdf') or '.pdf?' in url.lower() or '.pdf#' in url.lower():
                                link_text = ''.join(hyperlink.itertext())
                                preview = link_text[:40] + "..." if len(link_text) > 40 else link_text

                                issue = AccessibilityIssue(
                                    issue_type="PDF_LINK",
                                    description=f"Link to external PDF (verify PDF is accessible): \"{preview}\" -> {url[:60]}...",
                                    location=para_info.location,
                                    para_info=para_info
                                )
                                issues.append(issue)
                        except:
                            pass
            except:
                pass

        return issues

    def check_missing_toc(self) -> List[AccessibilityIssue]:
        """Detect missing table of contents in longer documents"""
        issues = []

        # Count paragraphs and check document length
        num_paragraphs = len(self.all_paragraphs)
        num_pages_estimate = num_paragraphs / 30  # Rough estimate: ~30 paragraphs per page

        # Check if document is "long" (>5 pages or >150 paragraphs)
        is_long_document = num_paragraphs > 150 or num_pages_estimate > 5

        if is_long_document:
            # Check if document has a table of contents
            has_toc = False

            # Method 1: Check for TOC field
            try:
                for para in self.target_doc.paragraphs:
                    # Look for TOC field code
                    para_element = para._element
                    fldChar_elements = para_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldChar')
                    if fldChar_elements:
                        # Check if this is a TOC field
                        instrText = para_element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}instrText')
                        if instrText is not None and 'TOC' in instrText.text:
                            has_toc = True
                            break
            except:
                pass

            # Method 2: Check for text that looks like a ToC heading
            if not has_toc:
                for para_info in self.all_paragraphs[:20]:  # Check first 20 paragraphs
                    para_text = para_info.paragraph.text.strip().lower()
                    if para_text in ['table of contents', 'contents', 'table of content']:
                        has_toc = True
                        break

            if not has_toc:
                issue = AccessibilityIssue(
                    issue_type="MISSING_TOC",
                    description=f"Long document (~{int(num_pages_estimate)} pages, {num_paragraphs} paragraphs) is missing a Table of Contents for navigation",
                    location="Document structure"
                )
                issues.append(issue)

        return issues

    def check_missing_bookmarks(self) -> List[AccessibilityIssue]:
        """Detect missing internal navigation/bookmarks in longer documents"""
        issues = []

        # Count internal bookmarks
        bookmark_count = 0
        internal_link_count = 0

        try:
            # Check for bookmarks in the document
            doc_element = self.target_doc.element
            bookmarks = doc_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bookmarkStart')
            bookmark_count = len(bookmarks)

            # Filter out built-in bookmarks (like _Toc, _Hlk, _GoBack)
            named_bookmarks = 0
            for bookmark in bookmarks:
                name = bookmark.get(qn('w:name'))
                if name and not name.startswith('_'):
                    named_bookmarks += 1

            # Check for internal hyperlinks (links to bookmarks)
            for para_info in self.all_paragraphs:
                para_element = para_info.paragraph._element
                hyperlinks = para_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink')

                for hyperlink in hyperlinks:
                    # Check if it's an internal link (has anchor attribute)
                    anchor = hyperlink.get(qn('w:anchor'))
                    if anchor:
                        internal_link_count += 1

        except:
            pass

        # Check if document is long and lacks bookmarks/internal navigation
        num_paragraphs = len(self.all_paragraphs)
        is_long_document = num_paragraphs > 150

        if is_long_document and named_bookmarks == 0 and internal_link_count == 0:
            issue = AccessibilityIssue(
                issue_type="MISSING_BOOKMARKS",
                description=f"Long document ({num_paragraphs} paragraphs) lacks internal bookmarks or navigation links (consider adding bookmarks for major sections)",
                location="Document structure"
            )
            issues.append(issue)

        return issues

    def check_nested_list_hierarchy(self) -> List[AccessibilityIssue]:
        """Flag incorrect or inconsistent nested list hierarchy"""
        issues = []

        # Track list levels and patterns
        prev_level = 0
        list_started = False

        for para_info in self.all_paragraphs:
            para = para_info.paragraph
            style_name = para.style.name

            # Check if this is a list paragraph
            if 'List' in style_name:
                list_started = True

                # Try to determine the level from the style name
                # Common patterns: "List Bullet", "List Bullet 2", "List Bullet 3", etc.
                # Or: "List Number", "List Number 2", etc.
                level = 1  # Default to level 1

                # Extract level from style name
                import re
                level_match = re.search(r'(\d+)$', style_name)
                if level_match:
                    level = int(level_match.group(1))

                # Check for level skipping (e.g., going from level 1 to level 3)
                if list_started and level > prev_level + 1:
                    issue = AccessibilityIssue(
                        issue_type="INCONSISTENT_LIST_HIERARCHY",
                        description=f"List skips from level {prev_level} to level {level} (should increment by 1). Text: \"{para.text[:50]}...\"",
                        location=para_info.location,
                        para_info=para_info
                    )
                    issues.append(issue)

                prev_level = level
            else:
                # Reset when we exit list context
                if list_started:
                    prev_level = 0
                    list_started = False

        return issues

    def check_layout_lists(self) -> List[AccessibilityIssue]:
        """Identify lists used for layout or indentation rather than semantic grouping"""
        issues = []

        # Track consecutive list items
        consecutive_lists = []
        current_list_group = []

        for para_info in self.all_paragraphs:
            para = para_info.paragraph
            style_name = para.style.name

            if 'List' in style_name:
                current_list_group.append(para_info)
            else:
                # End of list group
                if current_list_group:
                    consecutive_lists.append(current_list_group)
                    current_list_group = []

        # Add final group if exists
        if current_list_group:
            consecutive_lists.append(current_list_group)

        # Check each list group for potential layout misuse
        for list_group in consecutive_lists:
            # Check if list has only 1 item (likely used for indentation)
            if len(list_group) == 1:
                para_info = list_group[0]
                text = para_info.paragraph.text.strip()

                # If single list item is very long, it's likely layout abuse
                if len(text) > 200:
                    issue = AccessibilityIssue(
                        issue_type="LAYOUT_LIST",
                        description=f"Single-item list with long text (likely used for indentation): \"{text[:60]}...\"",
                        location=para_info.location,
                        para_info=para_info
                    )
                    issues.append(issue)

            # Check if all items are very dissimilar in length (indicates layout use)
            elif len(list_group) >= 3:
                lengths = [len(p.paragraph.text.strip()) for p in list_group]
                avg_length = sum(lengths) / len(lengths)

                # Calculate variance
                variance = sum((l - avg_length) ** 2 for l in lengths) / len(lengths)
                std_dev = variance ** 0.5

                # If high variance and some items are very long, likely layout
                if avg_length > 0 and std_dev / avg_length > 2 and max(lengths) > 300:
                    issue = AccessibilityIssue(
                        issue_type="LAYOUT_LIST",
                        description=f"List with inconsistent item lengths (may be used for layout). Items: {len(list_group)}",
                        location=list_group[0].location,
                        para_info=list_group[0]
                    )
                    issues.append(issue)

        return issues

    def check_long_sentences(self) -> List[AccessibilityIssue]:
        """Detect overly long sentences (>35 words)"""
        issues = []

        for para_info in self.all_paragraphs:
            para = para_info.paragraph
            text = para.text.strip()

            # Skip empty paragraphs
            if not text:
                continue

            # Split into sentences (simple split by . ! ?)
            # Use a more sophisticated sentence splitter
            import re
            sentences = re.split(r'[.!?]+\s+', text)

            for sentence in sentences:
                sentence = sentence.strip()
                if not sentence:
                    continue

                # Count words (split by whitespace)
                words = sentence.split()
                word_count = len(words)

                # Flag if more than 35 words
                if word_count > 35:
                    preview = sentence[:60] + "..." if len(sentence) > 60 else sentence

                    issue = AccessibilityIssue(
                        issue_type="LONG_SENTENCE",
                        description=f"Sentence has {word_count} words (>35 recommended): \"{preview}\"",
                        location=para_info.location,
                        para_info=para_info
                    )
                    issues.append(issue)
                    # Only flag first long sentence in paragraph to avoid overwhelming output
                    break

        return issues

    def check_excessive_formatting(self) -> List[AccessibilityIssue]:
        """Detect excessive or inconsistent use of bold, italics, or underline"""
        issues = []

        for para_info in self.all_paragraphs:
            para = para_info.paragraph
            text = para.text.strip()

            # Skip empty paragraphs
            if not text or len(text) < 10:
                continue

            # Skip headings (which are often bold)
            if 'Heading' in para.style.name:
                continue

            # Count characters with bold, italic, and underline
            total_chars = len(text)
            bold_chars = 0
            italic_chars = 0
            underline_chars = 0
            formatting_switches = 0
            prev_bold = False
            prev_italic = False
            prev_underline = False

            for run in para.runs:
                run_text = run.text
                run_len = len(run_text)

                if run.bold:
                    bold_chars += run_len
                if run.italic:
                    italic_chars += run_len
                if run.font.underline:
                    underline_chars += run_len

                # Track formatting switches (inconsistent use)
                if run.bold != prev_bold or run.italic != prev_italic or run.font.underline != prev_underline:
                    formatting_switches += 1

                prev_bold = run.bold
                prev_italic = run.italic
                prev_underline = run.font.underline

            # Flag if >50% of paragraph is bold, italic, or underlined
            if total_chars > 0:
                bold_ratio = bold_chars / total_chars
                italic_ratio = italic_chars / total_chars
                underline_ratio = underline_chars / total_chars

                # Flag excessive bold (>50%)
                if bold_ratio > 0.5:
                    preview = text[:60] + "..." if len(text) > 60 else text
                    issue = AccessibilityIssue(
                        issue_type="EXCESSIVE_BOLD",
                        description=f"Paragraph has {int(bold_ratio*100)}% bold text. If this is a heading, use a Heading style (Heading 1, 2, etc.) instead of bold. If intended for emphasis, use bold more selectively: \"{preview}\"",
                        location=para_info.location,
                        para_info=para_info
                    )
                    issues.append(issue)

                # Flag excessive italic (>50%)
                if italic_ratio > 0.5:
                    preview = text[:60] + "..." if len(text) > 60 else text
                    issue = AccessibilityIssue(
                        issue_type="EXCESSIVE_ITALIC",
                        description=f"Paragraph has {int(italic_ratio*100)}% italic text. If this is a heading or section title, use a Heading style instead. If intended for emphasis, use italics more selectively (large blocks of italic text are harder to read): \"{preview}\"",
                        location=para_info.location,
                        para_info=para_info
                    )
                    issues.append(issue)

                # Flag excessive underline (>30%) - lower threshold since underline should be rare
                if underline_ratio > 0.3:
                    preview = text[:60] + "..." if len(text) > 60 else text
                    issue = AccessibilityIssue(
                        issue_type="EXCESSIVE_UNDERLINE",
                        description=f"Paragraph has {int(underline_ratio*100)}% underlined text (reserve underline for links): \"{preview}\"",
                        location=para_info.location,
                        para_info=para_info
                    )
                    issues.append(issue)

                # Flag inconsistent formatting (many switches relative to paragraph length)
                # More than 8 formatting switches in a paragraph suggests inconsistent use
                if formatting_switches > 8 and len(para.runs) > 10:
                    preview = text[:60] + "..." if len(text) > 60 else text
                    issue = AccessibilityIssue(
                        issue_type="INCONSISTENT_FORMATTING",
                        description=f"Paragraph has inconsistent formatting ({formatting_switches} format changes): \"{preview}\"",
                        location=para_info.location,
                        para_info=para_info
                    )
                    issues.append(issue)

        return issues

    def check_images_alt_text(self) -> List[AccessibilityIssue]:
        """Detect images missing alt text"""
        issues = []

        # Find all inline shapes (images) in the document
        for para_info in self.all_paragraphs:
            para = para_info.paragraph

            # Check for inline shapes (images) in the paragraph
            try:
                para_element = para._element
                # Look for drawing elements (modern images)
                drawings = para_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')

                for drawing in drawings:
                    # Look for alt text in the docPr element
                    has_alt_text = False
                    docPr = drawing.find('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}docPr')

                    if docPr is not None:
                        # Check for title and description
                        title = docPr.get('title', '')
                        descr = docPr.get('descr', '')

                        if title or descr:
                            has_alt_text = True

                    if not has_alt_text:
                        issue = AccessibilityIssue(
                            issue_type="IMAGE_MISSING_ALT",
                            description="Image is missing alt text (right-click image > Edit Alt Text)",
                            location=para_info.location,
                            para_info=para_info
                        )
                        issues.append(issue)

                # Also check for older pict elements
                picts = para_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict')
                for pict in picts:
                    # Old format images often lack alt text
                    # Check for title attribute
                    has_alt_text = False
                    shape = pict.find('.//{urn:schemas-microsoft-com:vml}shape')
                    if shape is not None:
                        alt = shape.get('alt', '')
                        title = shape.get('title', '')
                        if alt or title:
                            has_alt_text = True

                    if not has_alt_text:
                        issue = AccessibilityIssue(
                            issue_type="IMAGE_MISSING_ALT",
                            description="Image is missing alt text (right-click image > Edit Alt Text)",
                            location=para_info.location,
                            para_info=para_info
                        )
                        issues.append(issue)

            except:
                pass  # If we can't parse images, skip

        return issues

    def check_decorative_images(self) -> List[AccessibilityIssue]:
        """Identify decorative images missing decorative marking"""
        issues = []

        # Find all inline shapes (images) in the document
        for para_info in self.all_paragraphs:
            para = para_info.paragraph

            try:
                para_element = para._element
                drawings = para_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')

                for drawing in drawings:
                    # Check if image has alt text
                    has_alt_text = False
                    is_marked_decorative = False

                    docPr = drawing.find('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}docPr')

                    if docPr is not None:
                        title = docPr.get('title', '')
                        descr = docPr.get('descr', '')

                        # Check if marked as decorative (empty title AND empty description)
                        # In Word, decorative images should have both set to empty string
                        if title == '' and descr == '':
                            is_marked_decorative = True
                        elif title or descr:
                            has_alt_text = True

                    # Heuristic: if image has no alt text and is not marked decorative,
                    # it might be decorative but not marked properly
                    # We can't definitively say it's decorative, so we'll flag for review
                    # Only flag if it doesn't have alt text (already covered by other check)
                    # This check is for images that ARE marked decorative but shouldn't be

                    # Actually, let's flag images that have empty alt text but might not be decorative
                    # Check if the image is likely meaningful by its context
                    if is_marked_decorative:
                        # Check surrounding text for context that suggests image is meaningful
                        para_text = para.text.strip()

                        # Keywords that suggest the image might be meaningful
                        meaningful_keywords = ['figure', 'diagram', 'chart', 'graph', 'screenshot',
                                             'illustration', 'photo', 'image', 'picture', 'logo']

                        # If paragraph mentions the image, it's probably not decorative
                        if any(keyword in para_text.lower() for keyword in meaningful_keywords):
                            issue = AccessibilityIssue(
                                issue_type="DECORATIVE_IMAGE_QUESTIONABLE",
                                description=f"Image marked as decorative but context suggests it may be meaningful: \"{para_text[:50]}...\"",
                                location=para_info.location,
                                para_info=para_info
                            )
                            issues.append(issue)

            except:
                pass

        return issues

    def check_image_text_content(self) -> List[AccessibilityIssue]:
        """Detect images that might contain text content (screenshots, schedules, tables)"""
        issues = []

        # Look for images with suspiciously large file sizes (likely screenshots)
        # or images in contexts that suggest they contain text
        for para_info in self.all_paragraphs:
            para = para_info.paragraph
            para_text = para.text.strip().lower()

            try:
                para_element = para._element
                drawings = para_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')

                for drawing in drawings:
                    # Check context - if nearby text mentions schedule, calendar, table, etc.
                    # this image might be a screenshot
                    context_keywords = ['schedule', 'calendar', 'table', 'dates', 'timeline',
                                       'grading scale', 'rubric', 'assignment', 'due date']

                    # Check alt text for hints
                    docPr = drawing.find('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}docPr')
                    alt_text = ""
                    if docPr is not None:
                        alt_text = docPr.get('title', '') + " " + docPr.get('descr', '')

                    # Check if image appears in a context suggesting it contains text
                    suspicious = False
                    reason = ""

                    # Check nearby text (current paragraph)
                    if any(keyword in para_text for keyword in context_keywords):
                        suspicious = True
                        reason = f"Image appears near text mentioning '{next(k for k in context_keywords if k in para_text)}'"

                    # Check alt text for keywords suggesting text content
                    alt_lower = alt_text.lower()
                    if any(keyword in alt_lower for keyword in ['screenshot', 'table', 'schedule', 'calendar']):
                        suspicious = True
                        reason = "Alt text suggests image contains structured data"

                    if suspicious:
                        issue = AccessibilityIssue(
                            issue_type="IMAGE_TEXT_CONTENT",
                            description=f"Image may contain text content that should be provided as actual text. {reason}. Consider using a real table or text list instead.",
                            location=para_info.location,
                            para_info=para_info
                        )
                        issues.append(issue)

            except:
                pass

        return issues

    def check_date_formats(self) -> List[AccessibilityIssue]:
        """Detect unclear date formats (numeric-only dates)"""
        issues = []

        import re

        # Patterns for numeric-only dates (ambiguous)
        # Examples: 12/5/2024, 5-12-24, 12.5.24
        numeric_date_pattern = r'\b\d{1,2}[-/.]\d{1,2}[-/.]\d{2,4}\b'

        for para_info in self.all_paragraphs:
            text = para_info.paragraph.text

            # Find all numeric date patterns
            matches = re.finditer(numeric_date_pattern, text)
            found_numeric_dates = []

            for match in matches:
                date_str = match.group()
                # Skip if it looks like a version number (e.g., 1.2.3)
                if not any(sep in date_str for sep in ['-', '/']):
                    continue
                found_numeric_dates.append(date_str)

            if found_numeric_dates:
                # Remove duplicates
                unique_dates = list(set(found_numeric_dates))
                if unique_dates:
                    example_dates = ', '.join(unique_dates[:3])
                    if len(unique_dates) > 3:
                        example_dates += f" (and {len(unique_dates) - 3} more)"

                    issue = AccessibilityIssue(
                        issue_type="NUMERIC_DATE_FORMAT",
                        description=f"Numeric date format found: {example_dates}. Use clear formats like 'January 15, 2024' or '15 Jan 2024' to avoid confusion.",
                        location=para_info.location,
                        para_info=para_info
                    )
                    issues.append(issue)

        return issues

    def check_merged_cells_layout(self) -> List[AccessibilityIssue]:
        """Identify merged table cells used to create visual layouts (already covered by check_table_merged_cells)"""
        # This is already implemented in check_table_merged_cells()
        # That function detects merged cells that may impact reading order
        # We'll enhance the existing function's description to emphasize layout issues
        return []

    def run_all_checks(self):
        """Run all accessibility checks and collect issues"""
        self.issues = []
        # Heading checks disabled - better suited for LLM analysis
        # self.issues.extend(self.check_unstyled_headings())
        # self.issues.extend(self.check_heading_hierarchy())

        # Algorithmic checks
        self.issues.extend(self.check_table_headers())
        self.issues.extend(self.check_layout_tables())
        self.issues.extend(self.check_empty_table_rows_columns())
        self.issues.extend(self.check_table_scope_declarations())
        self.issues.extend(self.check_table_captions())
        self.issues.extend(self.check_table_merged_cells())
        self.issues.extend(self.check_table_numeric_alignment())
        self.issues.extend(self.check_table_reading_order())
        self.issues.extend(self.check_table_color_coded_meaning())
        self.issues.extend(self.check_table_embedded_images())
        self.issues.extend(self.check_manual_alignment())
        self.issues.extend(self.check_font_sizes())
        self.issues.extend(self.check_decorative_fonts())
        self.issues.extend(self.check_inconsistent_fonts())
        self.issues.extend(self.check_color_contrast())
        self.issues.extend(self.check_color_only_meaning())
        self.issues.extend(self.check_text_over_backgrounds())
        self.issues.extend(self.check_document_metadata())
        self.issues.extend(self.check_document_language())
        self.issues.extend(self.check_multilingual_content())
        self.issues.extend(self.check_underline_non_links())
        self.issues.extend(self.check_line_spacing())
        self.issues.extend(self.check_full_justification())
        self.issues.extend(self.check_all_caps_blocks())
        self.issues.extend(self.check_non_descriptive_links())
        self.issues.extend(self.check_unstyled_links())
        self.issues.extend(self.check_long_urls())
        self.issues.extend(self.check_pdf_links())
        self.issues.extend(self.check_missing_toc())
        self.issues.extend(self.check_missing_bookmarks())
        self.issues.extend(self.check_nested_list_hierarchy())
        self.issues.extend(self.check_layout_lists())
        self.issues.extend(self.check_long_sentences())
        self.issues.extend(self.check_excessive_formatting())
        self.issues.extend(self.check_images_alt_text())
        self.issues.extend(self.check_decorative_images())
        self.issues.extend(self.check_image_text_content())
        self.issues.extend(self.check_date_formats())

    def create_marked_document(self, output_path: str):
        """Create a copy of the document with issues marked via comments and highlighting"""
        marked_doc = Document(self.target_path)

        # Group issues by location
        para_issues_by_location = {}
        table_issues = {}

        for issue in self.issues:
            if issue.para_info is not None:
                # For paragraph issues, group by table location if applicable
                key = (issue.para_info.table_idx, issue.para_info.row_idx, issue.para_info.cell_idx)
                if key not in para_issues_by_location:
                    para_issues_by_location[key] = []
                para_issues_by_location[key].append(issue)

            if issue.table_index is not None:
                if issue.table_index not in table_issues:
                    table_issues[issue.table_index] = []
                table_issues[issue.table_index].append(issue)

        # Mark top-level paragraphs (not in tables) - use inline markers like table cells
        for (table_idx, row_idx, cell_idx), issues_list in para_issues_by_location.items():
            if table_idx is None:
                # This is a top-level paragraph - find and mark it
                for issue in issues_list:
                    original_text = issue.para_info.paragraph.text.strip()

                    # Find the matching top-level paragraph
                    for para in marked_doc.paragraphs:
                        para_text = para.text.strip()

                        if para_text and (para_text in original_text or original_text in para_text):
                            # Add inline marker at the end
                            marker_run = para.add_run(f"  [← {issue.issue_type}]")
                            marker_run.font.color.rgb = RGBColor(255, 0, 0)
                            marker_run.bold = True
                            marker_run.font.size = Pt(9)
                            marker_run.font.highlight_color = WD_COLOR_INDEX.YELLOW

                            # Highlight the paragraph content
                            for run in para.runs[:-1]:  # Skip the marker we just added
                                run.font.highlight_color = WD_COLOR_INDEX.YELLOW

                            break

        # Mark paragraphs in table cells - mark each issue inline where it appears
        for (table_idx, row_idx, cell_idx), issues_list in para_issues_by_location.items():
            if table_idx is not None and table_idx < len(marked_doc.tables):
                table = marked_doc.tables[table_idx]
                if row_idx < len(table.rows) and cell_idx < len(table.rows[row_idx].cells):
                    cell = table.rows[row_idx].cells[cell_idx]

                    # For each issue, try to find and mark the specific text
                    for issue in issues_list:
                        # Get the original text from the issue
                        original_text = issue.para_info.paragraph.text.strip()

                        # Search through cell paragraphs for matching text
                        for para in cell.paragraphs:
                            para_text = para.text.strip()

                            # If this paragraph matches the issue text, mark it
                            if para_text and para_text in original_text or original_text in para_text:
                                # Add inline marker at the end of the paragraph
                                marker_run = para.add_run(f"  [← {issue.issue_type}]")
                                marker_run.font.color.rgb = RGBColor(255, 0, 0)
                                marker_run.bold = True
                                marker_run.font.size = Pt(9)
                                marker_run.font.highlight_color = WD_COLOR_INDEX.YELLOW

                                # Highlight the paragraph content
                                for run in para.runs[:-1]:  # Skip the marker we just added
                                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

                                break  # Found and marked this issue, move to next

        # Mark tables with issues - GROUP BY TABLE for conciseness
        for table_idx, issues_list in table_issues.items():
            if table_idx < len(marked_doc.tables):
                table = marked_doc.tables[table_idx]

                # Add marker in the first cell
                if table.rows and table.rows[0].cells:
                    first_cell = table.rows[0].cells[0]

                    # Group issues by type to avoid repetition
                    issue_summary = {}
                    for issue in issues_list:
                        issue_type = issue.issue_type
                        if issue_type not in issue_summary:
                            issue_summary[issue_type] = []
                        issue_summary[issue_type].append(issue.description)

                    # Create a single concise marker for all table issues
                    issue_count = len(issues_list)
                    issue_types_str = ", ".join(issue_summary.keys())

                    # Create summary text
                    if issue_count == 1:
                        summary_text = f"*** TABLE ISSUE: {issues_list[0].description} ***"
                    else:
                        summary_text = f"*** TABLE HAS {issue_count} ISSUES ({issue_types_str}) ***"
                        # Add details for unique issue types
                        for issue_type, descriptions in issue_summary.items():
                            # Only show first description for each type to avoid repetition
                            summary_text += f"\n  • {issue_type}: {descriptions[0]}"
                            if len(descriptions) > 1:
                                summary_text += f" (+{len(descriptions)-1} more)"

                    if first_cell.paragraphs:
                        first_para = first_cell.paragraphs[0]
                        marker_para = first_para.insert_paragraph_before(summary_text)
                    else:
                        marker_para = first_cell.add_paragraph(summary_text)

                    # Format the marker
                    for run in marker_para.runs:
                        run.font.color.rgb = RGBColor(255, 0, 0)
                        run.bold = True
                        run.font.size = Pt(11)
                        run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN

        marked_doc.save(output_path)

    def generate_report(self) -> str:
        """Generate comprehensive assessment report organized by category"""
        report_lines = []
        report_lines.append("=" * 80)
        report_lines.append("SYLLABUS ASSESSMENT REPORT")
        report_lines.append("=" * 80)
        report_lines.append("")

        # Document stats
        report_lines.append("DOCUMENT STATISTICS")
        report_lines.append("-" * 80)
        report_lines.append(f"Total paragraphs analyzed: {len(self.all_paragraphs)}")
        top_level = sum(1 for p in self.all_paragraphs if p.table_idx is None)
        in_tables = sum(1 for p in self.all_paragraphs if p.table_idx is not None)
        report_lines.append(f"  - Top-level paragraphs: {top_level}")
        report_lines.append(f"  - Paragraphs in tables: {in_tables}")
        report_lines.append(f"Total tables: {len(self.target_doc.tables)}")
        report_lines.append("")

        # Section completeness
        report_lines.append("SECTION COMPLETENESS")
        report_lines.append("-" * 80)
        missing = self.check_missing_sections()
        if missing:
            report_lines.append(f"[X] MISSING {len(missing)} REQUIRED SECTIONS:")
            for section in missing:
                report_lines.append(f"  - {section}")
        else:
            report_lines.append("[OK] All required sections present")
        report_lines.append("")

        # Heading structure
        report_lines.append("ACCESSIBILITY: HEADING STRUCTURE")
        report_lines.append("-" * 80)
        heading_check = self.check_heading_structure()
        if heading_check['counts']:
            report_lines.append("Heading usage:")
            for level, count in sorted(heading_check['counts'].items()):
                report_lines.append(f"  - {level}: {count}")
        if heading_check['issues']:
            for issue in heading_check['issues']:
                report_lines.append(f"[X] {issue}")
        else:
            report_lines.append("[OK] Heading structure looks good")
        report_lines.append("")

        # Heading level recommendations
        report_lines.append("ACCESSIBILITY: HEADING LEVEL RECOMMENDATIONS")
        report_lines.append("-" * 80)
        heading_recommendations = [issue for issue in self.issues
                    if issue.issue_type.startswith("SHOULD BE HEADING")]
        if heading_recommendations:
            report_lines.append(f"[X] Found {len(heading_recommendations)} headings that need to be corrected:")
            by_level = {}
            for issue in heading_recommendations:
                level = issue.issue_type
                if level not in by_level:
                    by_level[level] = []
                by_level[level].append(issue)
            for level in sorted(by_level.keys()):
                report_lines.append(f"\n  {level} ({len(by_level[level])} items):")
                for issue in by_level[level][:5]:
                    report_lines.append(f"    - {issue.location}: {issue.description}")
                if len(by_level[level]) > 5:
                    report_lines.append(f"    ... and {len(by_level[level]) - 5} more")
        else:
            report_lines.append("[OK] All headings are properly styled")
        report_lines.append("")

        # ============================================================
        # CATEGORY 1: FONT USAGE
        # ============================================================
        report_lines.append("=" * 80)
        report_lines.append("CATEGORY: FONT USAGE")
        report_lines.append("=" * 80)
        report_lines.append("")

        # Font sizes
        report_lines.append("ACCESSIBILITY: FONT SIZES")
        report_lines.append("-" * 80)
        font_issues = [issue for issue in self.issues if issue.issue_type == "SMALL_FONT"]
        if font_issues:
            report_lines.append(f"[X] Found {len(font_issues)} paragraph(s) with font size below 11pt:")
            report_lines.append("    (Minimum recommended size is 11pt for accessibility)")
            for issue in font_issues[:5]:
                report_lines.append(f"  - {issue.location}: {issue.description}")
            if len(font_issues) > 5:
                report_lines.append(f"  ... and {len(font_issues) - 5} more")
        else:
            report_lines.append("[OK] All text meets minimum font size requirements")
        report_lines.append("")

        # Decorative fonts
        report_lines.append("ACCESSIBILITY: DECORATIVE/INACCESSIBLE FONTS")
        report_lines.append("-" * 80)
        decorative_font_issues = [issue for issue in self.issues if issue.issue_type == "DECORATIVE_FONT"]
        if decorative_font_issues:
            report_lines.append(f"[X] Found {len(decorative_font_issues)} instance(s) of decorative/hard-to-read fonts:")
            report_lines.append("    (Use simple, readable fonts like Arial, Calibri, or Times New Roman)")
            for issue in decorative_font_issues[:5]:
                report_lines.append(f"  - {issue.location}: {issue.description}")
            if len(decorative_font_issues) > 5:
                report_lines.append(f"  ... and {len(decorative_font_issues) - 5} more")
        else:
            report_lines.append("[OK] All fonts are accessible and readable")
        report_lines.append("")

        # Font consistency
        report_lines.append("ACCESSIBILITY: FONT CONSISTENCY")
        report_lines.append("-" * 80)
        inconsistent_font_issues = [issue for issue in self.issues if issue.issue_type == "INCONSISTENT_FONTS"]
        if inconsistent_font_issues:
            for issue in inconsistent_font_issues:
                report_lines.append(f"[X] {issue.description}")
        else:
            report_lines.append("[OK] Font usage is consistent throughout the document")
        report_lines.append("")

        # ============================================================
        # CATEGORY 2: TABLE STRUCTURE
        # ============================================================
        report_lines.append("=" * 80)
        report_lines.append("CATEGORY: TABLE STRUCTURE")
        report_lines.append("=" * 80)
        report_lines.append("")

        # Table usage
        table_check = self.check_table_usage()
        report_lines.append("ACCESSIBILITY: TABLE USAGE")
        report_lines.append("-" * 80)
        report_lines.append(f"Total tables: {table_check['count']}")
        if table_check['issues']:
            for issue in table_check['issues']:
                report_lines.append(f"[X] {issue}")
        else:
            if table_check['count'] == 0:
                report_lines.append("[OK] No tables (good for accessibility)")
            else:
                report_lines.append("[OK] Table usage appears reasonable")
        report_lines.append("")

        # Empty table rows/columns
        report_lines.append("ACCESSIBILITY: EMPTY TABLE ROWS/COLUMNS")
        report_lines.append("-" * 80)
        empty_row_issues = [issue for issue in self.issues if issue.issue_type == "EMPTY_TABLE_ROW"]
        empty_col_issues = [issue for issue in self.issues if issue.issue_type == "EMPTY_TABLE_COLUMN"]
        if empty_row_issues or empty_col_issues:
            if empty_row_issues:
                report_lines.append(f"[X] Found {len(empty_row_issues)} empty table row(s) (used for spacing):")
                for issue in empty_row_issues:
                    report_lines.append(f"  - {issue.description}")
            if empty_col_issues:
                report_lines.append(f"[X] Found {len(empty_col_issues)} empty table column(s) (used for spacing):")
                for issue in empty_col_issues:
                    report_lines.append(f"  - {issue.description}")
            report_lines.append("    (Use proper table spacing/padding instead of empty rows/columns)")
        else:
            if len(self.target_doc.tables) > 0:
                report_lines.append("[OK] No empty table rows or columns detected")
            else:
                report_lines.append("[OK] No tables to check")
        report_lines.append("")

        # Layout tables
        report_lines.append("ACCESSIBILITY: LAYOUT vs. DATA TABLES")
        report_lines.append("-" * 80)
        layout_issues = [issue for issue in self.issues if issue.issue_type == "LAYOUT_TABLE"]
        if layout_issues:
            report_lines.append(f"[X] Found {len(layout_issues)} tables used for layout (should use headings/paragraphs):")
            for issue in layout_issues:
                report_lines.append(f"  - {issue.description}")
        else:
            if table_check['count'] > 0:
                report_lines.append("[OK] Tables appear to contain tabular data (not layout)")
            else:
                report_lines.append("[OK] No tables to check")
        report_lines.append("")

        # Table headers
        report_lines.append("ACCESSIBILITY: TABLE HEADERS")
        report_lines.append("-" * 80)
        table_header_issues = [issue for issue in self.issues if issue.issue_type == "TABLE_NO_HEADER"]
        if table_header_issues:
            report_lines.append(f"[X] Found {len(table_header_issues)} tables without proper headers:")
            for issue in table_header_issues:
                report_lines.append(f"  - {issue.description}")
        else:
            if table_check['count'] > 0:
                report_lines.append("[OK] All tables appear to have headers")
            else:
                report_lines.append("[OK] No tables to check")
        report_lines.append("")

        # Table scope declarations
        report_lines.append("ACCESSIBILITY: TABLE SCOPE DECLARATIONS")
        report_lines.append("-" * 80)
        scope_issues = [issue for issue in self.issues if issue.issue_type == "TABLE_MISSING_SCOPE"]
        if scope_issues:
            report_lines.append(f"[X] Found {len(scope_issues)} table(s) missing proper scope declarations:")
            for issue in scope_issues:
                report_lines.append(f"  - {issue.description}")
        else:
            if len(self.target_doc.tables) > 0:
                report_lines.append("[OK] Table scope declarations are in place")
            else:
                report_lines.append("[OK] No tables to check")
        report_lines.append("")

        # Table captions
        report_lines.append("ACCESSIBILITY: TABLE CAPTIONS/DESCRIPTIONS")
        report_lines.append("-" * 80)
        caption_issues = [issue for issue in self.issues if issue.issue_type == "TABLE_MISSING_CAPTION"]
        if caption_issues:
            report_lines.append(f"[X] Found {len(caption_issues)} table(s) missing captions or descriptions:")
            for issue in caption_issues:
                report_lines.append(f"  - {issue.description}")
        else:
            if len(self.target_doc.tables) > 0:
                report_lines.append("[OK] All tables have captions or descriptions")
            else:
                report_lines.append("[OK] No tables to check")
        report_lines.append("")

        # Table merged cells
        report_lines.append("ACCESSIBILITY: TABLE MERGED CELLS")
        report_lines.append("-" * 80)
        merged_cell_issues = [issue for issue in self.issues if issue.issue_type == "TABLE_MERGED_CELLS"]
        if merged_cell_issues:
            report_lines.append(f"[X] Found {len(merged_cell_issues)} table(s) with merged cells:")
            for issue in merged_cell_issues:
                report_lines.append(f"  - {issue.description}")
        else:
            if len(self.target_doc.tables) > 0:
                report_lines.append("[OK] No merged cells detected")
            else:
                report_lines.append("[OK] No tables to check")
        report_lines.append("")

        # Table numeric alignment
        report_lines.append("ACCESSIBILITY: TABLE NUMERIC ALIGNMENT")
        report_lines.append("-" * 80)
        numeric_align_issues = [issue for issue in self.issues if issue.issue_type == "TABLE_INCONSISTENT_NUMERIC_ALIGNMENT"]
        if numeric_align_issues:
            report_lines.append(f"[X] Found {len(numeric_align_issues)} table column(s) with inconsistent numeric alignment:")
            for issue in numeric_align_issues:
                report_lines.append(f"  - {issue.description}")
        else:
            if len(self.target_doc.tables) > 0:
                report_lines.append("[OK] Numeric data alignment is consistent")
            else:
                report_lines.append("[OK] No tables to check")
        report_lines.append("")

        # Table reading order
        report_lines.append("ACCESSIBILITY: TABLE READING ORDER")
        report_lines.append("-" * 80)
        reading_order_issues = [issue for issue in self.issues if issue.issue_type == "TABLE_COMPLEX_READING_ORDER"]
        if reading_order_issues:
            report_lines.append(f"[X] Found {len(reading_order_issues)} table(s) with complex reading order:")
            for issue in reading_order_issues:
                report_lines.append(f"  - {issue.description}")
        else:
            if len(self.target_doc.tables) > 0:
                report_lines.append("[OK] Table reading order appears logical")
            else:
                report_lines.append("[OK] No tables to check")
        report_lines.append("")

        # Table embedded images
        report_lines.append("ACCESSIBILITY: TABLE EMBEDDED IMAGES")
        report_lines.append("-" * 80)
        embedded_image_issues = [issue for issue in self.issues if issue.issue_type == "TABLE_EMBEDDED_IMAGES"]
        if embedded_image_issues:
            report_lines.append(f"[X] Found {len(embedded_image_issues)} table(s) with embedded images:")
            for issue in embedded_image_issues:
                report_lines.append(f"  - {issue.description}")
        else:
            if len(self.target_doc.tables) > 0:
                report_lines.append("[OK] No embedded images in table cells")
            else:
                report_lines.append("[OK] No tables to check")
        report_lines.append("")

        # ============================================================
        # CATEGORY 3: COLOR & CONTRAST
        # ============================================================
        report_lines.append("=" * 80)
        report_lines.append("CATEGORY: COLOR & CONTRAST")
        report_lines.append("=" * 80)
        report_lines.append("")

        # Color contrast
        report_lines.append("ACCESSIBILITY: COLOR CONTRAST")
        report_lines.append("-" * 80)
        contrast_issues = [issue for issue in self.issues if issue.issue_type == "LOW_CONTRAST"]
        if contrast_issues:
            report_lines.append(f"[X] Found {len(contrast_issues)} instance(s) of insufficient color contrast:")
            report_lines.append("    (WCAG AA requires 4.5:1 for normal text, 3:1 for large text)")
            for issue in contrast_issues[:5]:
                report_lines.append(f"  - {issue.location}: {issue.description}")
            if len(contrast_issues) > 5:
                report_lines.append(f"  ... and {len(contrast_issues) - 5} more")
        else:
            report_lines.append("[OK] All text meets WCAG AA color contrast requirements")
        report_lines.append("")

        # Color as sole indicator
        report_lines.append("ACCESSIBILITY: COLOR AS SOLE INDICATOR")
        report_lines.append("-" * 80)
        color_only_issues = [issue for issue in self.issues if issue.issue_type == "COLOR_ONLY_MEANING"]
        color_coded_table_issues = [issue for issue in self.issues if issue.issue_type == "COLOR_CODED_TABLE"]
        if color_only_issues or color_coded_table_issues:
            if color_only_issues:
                report_lines.append(f"[X] Found {len(color_only_issues)} instance(s) where color may be the only indicator:")
                report_lines.append("    (Add bold, italic, or other formatting to supplement color)")
                for issue in color_only_issues[:5]:
                    report_lines.append(f"  - {issue.location}: {issue.description}")
                if len(color_only_issues) > 5:
                    report_lines.append(f"  ... and {len(color_only_issues) - 5} more")
            if color_coded_table_issues:
                report_lines.append(f"[X] Found {len(color_coded_table_issues)} color-coded table(s):")
                for issue in color_coded_table_issues:
                    report_lines.append(f"  - {issue.description}")
        else:
            report_lines.append("[OK] No issues with color as sole indicator detected")
        report_lines.append("")

        # Text over backgrounds
        report_lines.append("ACCESSIBILITY: TEXT OVER COLORED BACKGROUNDS")
        report_lines.append("-" * 80)
        text_bg_issues = [issue for issue in self.issues if issue.issue_type == "TEXT_OVER_BACKGROUND"]
        if text_bg_issues:
            report_lines.append(f"[X] Found {len(text_bg_issues)} instance(s) of text over colored backgrounds with low contrast:")
            report_lines.append("    (Ensure sufficient contrast between text and background colors)")
            for issue in text_bg_issues[:5]:
                report_lines.append(f"  - {issue.location}: {issue.description}")
            if len(text_bg_issues) > 5:
                report_lines.append(f"  ... and {len(text_bg_issues) - 5} more")
        else:
            report_lines.append("[OK] No issues with text over colored backgrounds detected")
        report_lines.append("")

        # ============================================================
        # CATEGORY 4: LINKS & NAVIGATION
        # ============================================================
        report_lines.append("=" * 80)
        report_lines.append("CATEGORY: LINKS & NAVIGATION")
        report_lines.append("=" * 80)
        report_lines.append("")

        # Non-descriptive links
        report_lines.append("ACCESSIBILITY: NON-DESCRIPTIVE LINKS")
        report_lines.append("-" * 80)
        non_desc_link_issues = [issue for issue in self.issues if issue.issue_type == "NON_DESCRIPTIVE_LINK"]
        if non_desc_link_issues:
            report_lines.append(f"[X] Found {len(non_desc_link_issues)} link(s) with non-descriptive text:")
            report_lines.append("    (Use descriptive text like 'Course Syllabus' instead of 'click here')")
            for issue in non_desc_link_issues[:5]:
                report_lines.append(f"  - {issue.location}: {issue.description}")
            if len(non_desc_link_issues) > 5:
                report_lines.append(f"  ... and {len(non_desc_link_issues) - 5} more")
        else:
            report_lines.append("[OK] No non-descriptive links detected")
        report_lines.append("")

        # Unstyled links
        report_lines.append("ACCESSIBILITY: UNSTYLED LINKS")
        report_lines.append("-" * 80)
        unstyled_link_issues = [issue for issue in self.issues if issue.issue_type == "UNSTYLED_LINK"]
        if unstyled_link_issues:
            report_lines.append(f"[X] Found {len(unstyled_link_issues)} link(s) styled as normal text:")
            report_lines.append("    (Links should be underlined and/or colored to be distinguishable)")
            for issue in unstyled_link_issues[:5]:
                report_lines.append(f"  - {issue.location}: {issue.description}")
            if len(unstyled_link_issues) > 5:
                report_lines.append(f"  ... and {len(unstyled_link_issues) - 5} more")
        else:
            report_lines.append("[OK] All links have distinguishable styling")
        report_lines.append("")

        # Long URLs
        report_lines.append("ACCESSIBILITY: LONG URLs")
        report_lines.append("-" * 80)
        long_url_issues = [issue for issue in self.issues if issue.issue_type == "LONG_URL"]
        if long_url_issues:
            report_lines.append(f"[X] Found {len(long_url_issues)} excessively long URL(s):")
            report_lines.append("    (Consider using descriptive link text instead of displaying long URLs)")
            for issue in long_url_issues[:5]:
                report_lines.append(f"  - {issue.location}: {issue.description}")
            if len(long_url_issues) > 5:
                report_lines.append(f"  ... and {len(long_url_issues) - 5} more")
        else:
            report_lines.append("[OK] No excessively long URLs detected")
        report_lines.append("")

        # PDF links
        report_lines.append("ACCESSIBILITY: LINKS TO PDFs")
        report_lines.append("-" * 80)
        pdf_link_issues = [issue for issue in self.issues if issue.issue_type == "PDF_LINK"]
        if pdf_link_issues:
            report_lines.append(f"[X] Found {len(pdf_link_issues)} link(s) to external PDF files:")
            report_lines.append("    (Verify that linked PDFs are accessible and properly tagged)")
            for issue in pdf_link_issues[:5]:
                report_lines.append(f"  - {issue.location}: {issue.description}")
            if len(pdf_link_issues) > 5:
                report_lines.append(f"  ... and {len(pdf_link_issues) - 5} more")
        else:
            report_lines.append("[OK] No links to external PDFs detected")
        report_lines.append("")

        # Table of contents
        report_lines.append("ACCESSIBILITY: TABLE OF CONTENTS")
        report_lines.append("-" * 80)
        toc_issues = [issue for issue in self.issues if issue.issue_type == "MISSING_TOC"]
        if toc_issues:
            for issue in toc_issues:
                report_lines.append(f"[X] {issue.description}")
            report_lines.append("    (Long documents benefit from a table of contents for navigation)")
        else:
            report_lines.append("[OK] Document has table of contents or is short enough not to require one")
        report_lines.append("")

        # Bookmarks
        report_lines.append("ACCESSIBILITY: INTERNAL NAVIGATION/BOOKMARKS")
        report_lines.append("-" * 80)
        bookmark_issues = [issue for issue in self.issues if issue.issue_type == "MISSING_BOOKMARKS"]
        if bookmark_issues:
            for issue in bookmark_issues:
                report_lines.append(f"[X] {issue.description}")
            report_lines.append("    (Add bookmarks to major sections for easier navigation)")
        else:
            report_lines.append("[OK] Document has internal navigation or is short enough not to require it")
        report_lines.append("")

        # ============================================================
        # CATEGORY 5: LISTS
        # ============================================================
        report_lines.append("=" * 80)
        report_lines.append("CATEGORY: LISTS")
        report_lines.append("=" * 80)
        report_lines.append("")

        # List formatting
        report_lines.append("ACCESSIBILITY: LIST FORMATTING")
        report_lines.append("-" * 80)
        list_check = self.check_list_usage()
        report_lines.append(f"Proper list items: {list_check['list_count']}")
        report_lines.append(f"Manual list items: {list_check['manual_list_count']}")
        if list_check['issues']:
            for issue in list_check['issues']:
                report_lines.append(f"[X] {issue}")
        else:
            report_lines.append("[OK] List formatting looks good")
        report_lines.append("")

        # Nested list hierarchy
        report_lines.append("ACCESSIBILITY: NESTED LIST HIERARCHY")
        report_lines.append("-" * 80)
        list_hierarchy_issues = [issue for issue in self.issues if issue.issue_type == "INCONSISTENT_LIST_HIERARCHY"]
        if list_hierarchy_issues:
            report_lines.append(f"[X] Found {len(list_hierarchy_issues)} inconsistent list level(s):")
            report_lines.append("    (List levels should increment by 1, not skip levels)")
            for issue in list_hierarchy_issues[:5]:
                report_lines.append(f"  - {issue.location}: {issue.description}")
            if len(list_hierarchy_issues) > 5:
                report_lines.append(f"  ... and {len(list_hierarchy_issues) - 5} more")
        else:
            report_lines.append("[OK] List hierarchy is consistent")
        report_lines.append("")

        # Layout lists
        report_lines.append("ACCESSIBILITY: LISTS USED FOR LAYOUT")
        report_lines.append("-" * 80)
        layout_list_issues = [issue for issue in self.issues if issue.issue_type == "LAYOUT_LIST"]
        if layout_list_issues:
            report_lines.append(f"[X] Found {len(layout_list_issues)} list(s) potentially used for layout/indentation:")
            report_lines.append("    (Lists should be used for semantic grouping, not layout)")
            for issue in layout_list_issues[:5]:
                report_lines.append(f"  - {issue.location}: {issue.description}")
            if len(layout_list_issues) > 5:
                report_lines.append(f"  ... and {len(layout_list_issues) - 5} more")
        else:
            report_lines.append("[OK] Lists appear to be used semantically")
        report_lines.append("")

        # ============================================================
        # CATEGORY 6: TEXT FORMATTING
        # ============================================================
        report_lines.append("=" * 80)
        report_lines.append("CATEGORY: TEXT FORMATTING")
        report_lines.append("=" * 80)
        report_lines.append("")

        # Manual alignment (pseudo-tables)
        report_lines.append("ACCESSIBILITY: MANUAL ALIGNMENT (PSEUDO-TABLES)")
        report_lines.append("-" * 80)
        pseudo_table_issues = [issue for issue in self.issues if issue.issue_type == "PSEUDO_TABLE"]
        if pseudo_table_issues:
            report_lines.append(f"[X] Found {len(pseudo_table_issues)} paragraph(s) using tabs or spaces for alignment:")
            report_lines.append("    (This creates pseudo-tables that are difficult for screen readers to parse)")
            for issue in pseudo_table_issues[:5]:
                report_lines.append(f"  - {issue.location}: {issue.description}")
            if len(pseudo_table_issues) > 5:
                report_lines.append(f"  ... and {len(pseudo_table_issues) - 5} more")
        else:
            report_lines.append("[OK] No manual alignment detected")
        report_lines.append("")

        # Underlined text
        report_lines.append("ACCESSIBILITY: UNDERLINED TEXT")
        report_lines.append("-" * 80)
        underline_issues = [issue for issue in self.issues if issue.issue_type == "UNDERLINE_NON_LINK"]
        if underline_issues:
            report_lines.append(f"[X] Found {len(underline_issues)} underlined text(s) that are not hyperlinks:")
            report_lines.append("    (Underlined text is often confused with links by users)")
            for issue in underline_issues[:5]:
                report_lines.append(f"  - {issue.location}: {issue.description}")
            if len(underline_issues) > 5:
                report_lines.append(f"  ... and {len(underline_issues) - 5} more")
        else:
            report_lines.append("[OK] No non-link underlined text detected")
        report_lines.append("")

        # Line spacing
        report_lines.append("ACCESSIBILITY: LINE SPACING")
        report_lines.append("-" * 80)
        spacing_issues = [issue for issue in self.issues if issue.issue_type == "LOW_LINE_SPACING"]
        if spacing_issues:
            report_lines.append(f"[X] Found {len(spacing_issues)} paragraph(s) with line spacing below 1.15:")
            report_lines.append("    (WCAG recommends minimum 1.15 for readability)")
            for issue in spacing_issues[:5]:
                report_lines.append(f"  - {issue.location}: {issue.description}")
            if len(spacing_issues) > 5:
                report_lines.append(f"  ... and {len(spacing_issues) - 5} more")
        else:
            report_lines.append("[OK] Line spacing meets minimum requirements")
        report_lines.append("")

        # Text justification
        report_lines.append("ACCESSIBILITY: TEXT JUSTIFICATION")
        report_lines.append("-" * 80)
        justify_issues = [issue for issue in self.issues if issue.issue_type == "FULL_JUSTIFICATION"]
        if justify_issues:
            report_lines.append(f"[X] Found {len(justify_issues)} paragraph(s) with full justification:")
            report_lines.append("    (Creates uneven spacing, harder to read especially for dyslexic readers)")
            for issue in justify_issues[:5]:
                report_lines.append(f"  - {issue.location}: {issue.description}")
            if len(justify_issues) > 5:
                report_lines.append(f"  ... and {len(justify_issues) - 5} more")
        else:
            report_lines.append("[OK] No full justification detected")
        report_lines.append("")

        # ALL CAPS blocks
        report_lines.append("ACCESSIBILITY: ALL CAPS BLOCKS")
        report_lines.append("-" * 80)
        caps_issues = [issue for issue in self.issues if issue.issue_type == "ALL_CAPS_BLOCK"]
        if caps_issues:
            report_lines.append(f"[X] Found {len(caps_issues)} large block(s) of ALL CAPS text:")
            report_lines.append("    (ALL CAPS text is harder to read than mixed case)")
            for issue in caps_issues[:5]:
                report_lines.append(f"  - {issue.location}: {issue.description}")
            if len(caps_issues) > 5:
                report_lines.append(f"  ... and {len(caps_issues) - 5} more")
        else:
            report_lines.append("[OK] No large ALL CAPS blocks detected")
        report_lines.append("")

        # Excessive formatting
        report_lines.append("ACCESSIBILITY: EXCESSIVE/INCONSISTENT FORMATTING")
        report_lines.append("-" * 80)
        excessive_bold = [issue for issue in self.issues if issue.issue_type == "EXCESSIVE_BOLD"]
        excessive_italic = [issue for issue in self.issues if issue.issue_type == "EXCESSIVE_ITALIC"]
        excessive_underline = [issue for issue in self.issues if issue.issue_type == "EXCESSIVE_UNDERLINE"]
        inconsistent_formatting = [issue for issue in self.issues if issue.issue_type == "INCONSISTENT_FORMATTING"]
        all_formatting_issues = excessive_bold + excessive_italic + excessive_underline + inconsistent_formatting
        if all_formatting_issues:
            if excessive_bold:
                report_lines.append(f"[X] Found {len(excessive_bold)} paragraph(s) with excessive bold:")
                for issue in excessive_bold[:3]:
                    report_lines.append(f"  - {issue.location}: {issue.description}")
            if excessive_italic:
                report_lines.append(f"[X] Found {len(excessive_italic)} paragraph(s) with excessive italics:")
                for issue in excessive_italic[:3]:
                    report_lines.append(f"  - {issue.location}: {issue.description}")
            if excessive_underline:
                report_lines.append(f"[X] Found {len(excessive_underline)} paragraph(s) with excessive underline:")
                for issue in excessive_underline[:3]:
                    report_lines.append(f"  - {issue.location}: {issue.description}")
            if inconsistent_formatting:
                report_lines.append(f"[X] Found {len(inconsistent_formatting)} paragraph(s) with inconsistent formatting:")
                for issue in inconsistent_formatting[:3]:
                    report_lines.append(f"  - {issue.location}: {issue.description}")
            report_lines.append("    (Excessive or inconsistent formatting reduces effectiveness and readability)")
        else:
            report_lines.append("[OK] Formatting usage is appropriate")
        report_lines.append("")

        # ============================================================
        # CATEGORY 7: READABILITY
        # ============================================================
        report_lines.append("=" * 80)
        report_lines.append("CATEGORY: READABILITY")
        report_lines.append("=" * 80)
        report_lines.append("")

        # Long sentences
        report_lines.append("ACCESSIBILITY: SENTENCE LENGTH")
        report_lines.append("-" * 80)
        long_sentence_issues = [issue for issue in self.issues if issue.issue_type == "LONG_SENTENCE"]
        if long_sentence_issues:
            report_lines.append(f"[X] Found {len(long_sentence_issues)} overly long sentence(s) (>35 words):")
            report_lines.append("    (Shorter sentences improve readability and comprehension)")
            for issue in long_sentence_issues[:5]:
                report_lines.append(f"  - {issue.location}: {issue.description}")
            if len(long_sentence_issues) > 5:
                report_lines.append(f"  ... and {len(long_sentence_issues) - 5} more")
        else:
            report_lines.append("[OK] Sentence lengths are within recommended limits")
        report_lines.append("")

        # Date formats
        report_lines.append("ACCESSIBILITY: DATE FORMATS")
        report_lines.append("-" * 80)
        date_format_issues = [issue for issue in self.issues if issue.issue_type == "NUMERIC_DATE_FORMAT"]
        if date_format_issues:
            report_lines.append(f"[X] Found {len(date_format_issues)} paragraph(s) with numeric-only date formats:")
            report_lines.append("    (Numeric dates like 12/5/24 are ambiguous - use 'December 5, 2024' or '5 Dec 2024')")
            for issue in date_format_issues[:5]:
                report_lines.append(f"  - {issue.location}: {issue.description}")
            if len(date_format_issues) > 5:
                report_lines.append(f"  ... and {len(date_format_issues) - 5} more")
        else:
            report_lines.append("[OK] No ambiguous numeric date formats detected")
        report_lines.append("")

        # ============================================================
        # CATEGORY 8: IMAGES
        # ============================================================
        report_lines.append("=" * 80)
        report_lines.append("CATEGORY: IMAGES")
        report_lines.append("=" * 80)
        report_lines.append("")

        # Image alt text
        report_lines.append("ACCESSIBILITY: IMAGE ALT TEXT")
        report_lines.append("-" * 80)
        image_alt_issues = [issue for issue in self.issues if issue.issue_type == "IMAGE_MISSING_ALT"]
        if image_alt_issues:
            report_lines.append(f"[X] Found {len(image_alt_issues)} image(s) missing alt text:")
            report_lines.append("    (All images need alt text or must be marked as decorative)")
            for issue in image_alt_issues[:5]:
                report_lines.append(f"  - {issue.location}: {issue.description}")
            if len(image_alt_issues) > 5:
                report_lines.append(f"  ... and {len(image_alt_issues) - 5} more")
        else:
            report_lines.append("[OK] All images have alt text or no images detected")
        report_lines.append("")

        # Decorative images
        report_lines.append("ACCESSIBILITY: DECORATIVE IMAGE MARKING")
        report_lines.append("-" * 80)
        decorative_issues = [issue for issue in self.issues if issue.issue_type == "DECORATIVE_IMAGE_QUESTIONABLE"]
        if decorative_issues:
            report_lines.append(f"[X] Found {len(decorative_issues)} image(s) marked as decorative but may be meaningful:")
            report_lines.append("    (Verify that decorative marking is appropriate for these images)")
            for issue in decorative_issues[:5]:
                report_lines.append(f"  - {issue.location}: {issue.description}")
            if len(decorative_issues) > 5:
                report_lines.append(f"  ... and {len(decorative_issues) - 5} more")
        else:
            report_lines.append("[OK] Decorative image marking appears appropriate")
        report_lines.append("")

        # Images containing text content
        report_lines.append("ACCESSIBILITY: IMAGES CONTAINING TEXT/SCHEDULES")
        report_lines.append("-" * 80)
        image_text_issues = [issue for issue in self.issues if issue.issue_type == "IMAGE_TEXT_CONTENT"]
        if image_text_issues:
            report_lines.append(f"[X] Found {len(image_text_issues)} image(s) that may contain text content:")
            report_lines.append("    (Text in images is inaccessible to screen readers - use actual text/tables instead)")
            for issue in image_text_issues[:5]:
                report_lines.append(f"  - {issue.location}: {issue.description}")
            if len(image_text_issues) > 5:
                report_lines.append(f"  ... and {len(image_text_issues) - 5} more")
        else:
            report_lines.append("[OK] No images appear to contain text content")
        report_lines.append("")

        # General images info
        report_lines.append("ACCESSIBILITY: IMAGES")
        report_lines.append("-" * 80)
        image_check = self.check_images()
        if image_check['issues']:
            for issue in image_check['issues']:
                report_lines.append(f"[i] {issue}")
        else:
            report_lines.append("[OK] No images detected")
        report_lines.append("")

        # ============================================================
        # CATEGORY 9: DOCUMENT PROPERTIES
        # ============================================================
        report_lines.append("=" * 80)
        report_lines.append("CATEGORY: DOCUMENT PROPERTIES")
        report_lines.append("=" * 80)
        report_lines.append("")

        # Document metadata
        report_lines.append("ACCESSIBILITY: DOCUMENT METADATA")
        report_lines.append("-" * 80)
        metadata_issues = [issue for issue in self.issues if issue.issue_type == "MISSING_TITLE"]
        if metadata_issues:
            for issue in metadata_issues:
                report_lines.append(f"[X] {issue.description}")
        else:
            report_lines.append("[OK] Document has title metadata")
        report_lines.append("")

        # Document language
        report_lines.append("ACCESSIBILITY: DOCUMENT LANGUAGE SETTING")
        report_lines.append("-" * 80)
        language_issues = [issue for issue in self.issues if issue.issue_type in ["MISSING_LANGUAGE", "INVALID_LANGUAGE"]]
        if language_issues:
            for issue in language_issues:
                report_lines.append(f"[X] {issue.description}")
        else:
            report_lines.append("[OK] Document has valid language setting")
        report_lines.append("")

        # Multilingual content
        report_lines.append("ACCESSIBILITY: MULTILINGUAL CONTENT")
        report_lines.append("-" * 80)
        multilingual_issues = [issue for issue in self.issues if issue.issue_type == "MULTIPLE_LANGUAGES"]
        untagged_language_issues = [issue for issue in self.issues if issue.issue_type == "UNTAGGED_LANGUAGE"]
        if multilingual_issues or untagged_language_issues:
            if multilingual_issues:
                for issue in multilingual_issues:
                    report_lines.append(f"[i] {issue.description}")
            if untagged_language_issues:
                report_lines.append(f"[X] Found potential multilingual content without proper language tagging:")
                for issue in untagged_language_issues:
                    report_lines.append(f"  - {issue.description}")
        else:
            report_lines.append("[OK] No multilingual content issues detected")
        report_lines.append("")

        # ============================================================
        # SUMMARY
        # ============================================================
        report_lines.append("=" * 80)
        report_lines.append("SUMMARY")
        report_lines.append("=" * 80)

        total_issues = (
            len(missing) +
            len(heading_check['issues']) +
            len(table_check['issues']) +
            len(table_header_issues) +
            len(layout_issues) +
            len(list_check['issues']) +
            len(pseudo_table_issues) +
            len(font_issues) +
            len(decorative_font_issues) +
            len(inconsistent_font_issues) +
            len(contrast_issues) +
            len(color_only_issues) +
            len(color_coded_table_issues) +
            len(text_bg_issues) +
            len(empty_row_issues) +
            len(empty_col_issues) +
            len(metadata_issues) +
            len(language_issues) +
            len(untagged_language_issues) +
            len(underline_issues) +
            len(spacing_issues) +
            len(justify_issues) +
            len(caps_issues) +
            len(non_desc_link_issues) +
            len(long_sentence_issues) +
            len(date_format_issues) +
            len(image_alt_issues) +
            len(decorative_issues) +
            len(image_text_issues)
        )

        if total_issues == 0:
            report_lines.append("[OK] Excellent! This syllabus meets all standards.")
        elif total_issues < 5:
            report_lines.append(f"[!] This syllabus has {total_issues} issues that should be addressed.")
        else:
            report_lines.append(f"[X] This syllabus has {total_issues} issues that need attention.")

        report_lines.append("")
        report_lines.append("RECOMMENDATIONS:")
        report_lines.append("1. Add all missing required sections from the template")
        report_lines.append("2. Use Heading 1, Heading 2, etc. styles (not bold text) for section headers")
        report_lines.append("3. Avoid using tables for layout; use headings and paragraphs instead")
        report_lines.append("4. Ensure all data tables have proper header rows")
        report_lines.append("5. Use descriptive hyperlink text instead of displaying URLs")
        report_lines.append("6. Ensure all images have descriptive alt text")
        report_lines.append("7. Use built-in list formatting instead of manual bullets")
        report_lines.append("8. Don't use tabs or multiple spaces for alignment; use tables or proper formatting")
        report_lines.append("")

        return "\n".join(report_lines)


def main():
    """Main entry point for command-line usage"""
    import os

    parser = argparse.ArgumentParser(
        description='Check a syllabus against UWM template standards',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  %(prog)s my_syllabus.docx
  %(prog)s english215.docx
        """
    )
    parser.add_argument('syllabus', help='Path to the syllabus to check (DOCX)')
    parser.add_argument('-t', '--template',
                       default='Uniform-Syllabus-Template-1.docx',
                       help='Path to template (default: Uniform-Syllabus-Template-1.docx)')

    args = parser.parse_args()

    # Auto-generate output file names based on input syllabus
    base_name = os.path.splitext(args.syllabus)[0]
    output_report = f"{base_name}_report.txt"
    output_marked = f"{base_name}_marked.docx"

    print(f"Analyzing: {args.syllabus}")
    print(f"Template: {args.template}")
    print()

    checker = SyllabusChecker(args.template, args.syllabus)

    # Show document stats
    print(f"Found {len(checker.all_paragraphs)} total paragraphs")
    top_level = sum(1 for p in checker.all_paragraphs if p.table_idx is None)
    in_tables = sum(1 for p in checker.all_paragraphs if p.table_idx is not None)
    print(f"  - {top_level} top-level paragraphs")
    print(f"  - {in_tables} paragraphs in tables")
    print()

    print("Running accessibility checks...")
    checker.run_all_checks()

    # Generate report
    report = checker.generate_report()

    # Always save text report
    with open(output_report, 'w') as f:
        f.write(report)
    print(f"✓ Text report saved to: {output_report}")

    # Always create marked document
    print(f"Creating marked-up document...")
    checker.create_marked_document(output_marked)
    print(f"✓ Marked document saved to: {output_marked}")

    print()
    print("=" * 60)
    print("DONE! Review the following files:")
    print(f"  1. {output_report} - Detailed text report")
    print(f"  2. {output_marked} - Word doc with issues highlighted")
    print("=" * 60)


if __name__ == '__main__':
    main()
